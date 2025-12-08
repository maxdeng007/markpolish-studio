import streamlit as st
import markdown
import re
import urllib.parse
import time
import os
from datetime import datetime

# --- UI TRANSLATIONS ---
TRANSLATIONS = {
    "en": {
        "app_title": "MarkPolish V1.0",
        "app_subtitle": "Content Ops Edition",
        "simple_mode": "Simple Mode",
        "dark_mode": "Dark Mode",
        "language": "Language",
        "ai_assistant": "AI Assistant",
        "files_templates": "Files & Templates",
        "add_components": "Add Components",
        "appearance": "Appearance",
        "plugins": "Plugins",
        "preview_mode": "Preview Mode",
        "editor": "Editor",
        "preview": "Preview",
        "mobile": "Mobile",
        "pc": "PC",
        "words": "words",
        "chars": "chars",
        "min_read": "min read",
        "saved": "Saved",
        "saving": "Saving...",
        "save_error": "Save error",
        "quick_save": "Quick Save",
        "undo": "Undo",
        "redo": "Redo",
        "export": "Export",
        "copy_wechat": "Copy for WeChat",
        "standard_html": "Standard HTML",
        "download": "Download",
        "theme": "Theme",
        "image_source": "Image Source",
        "files": "Files",
        "templates": "Templates",
        "no_content": "No content to preview. Start typing in the editor!",
        "layout": "Layout",
        "content": "Content",
        "interactive": "Interactive",
        "media": "Media",
        "comp_hero": "Hero",
        "comp_2col": "2-Col",
        "comp_3col": "3-Col",
        "comp_steps": "Steps",
        "comp_timeline": "Timeline",
        "comp_table": "Table",
        "comp_card": "Card",
        "comp_reveal": "Reveal",
        "comp_badge": "Badge",
        "comp_button": "Button",
        "comp_image": "AI Image",
        "files": "Files",
        "templates": "Templates",
        "save": "Save",
        "load": "Load",
        "delete": "Delete",
        "version_history": "Version History",
        "category": "Category",
        "all_templates": "All Templates",
        "use_template": "Use Template",
        "preview": "Preview",
        "search_templates": "Search templates",
        "no_files": "No files yet",
        "file_name": "File Name",
        "click_to_insert": "Click to insert components",
        "markdown_help": "Markdown Syntax Help",
        "image_assets": "Image & Assets",
        "upload_image": "Upload Image",
        "image_library": "Image Library",
        "no_images": "No images yet",
        "insert_image": "Insert",
        "preview_mode": "Preview Mode",
        "mobile_preview": "Mobile",
        "pc_preview": "PC",
        "reload_plugins": "Reload",
        "create_plugin": "Create Plugin",
        "installed_plugins": "Installed Plugins",
        "no_plugins": "No plugins installed",
        "no_files_search": "No files match your search.",
        "no_files_yet": "No saved files yet. Create a new file or use a template to get started.",
        "no_templates_match": "No templates match your search.",
        "click_to_insert_hint": "Click any button to insert. ✅=WeChat compatible, 🌐=HTML only",
        "view_full_guide": "View Full Guide",
        "syntax_tip": "Tip: Click any syntax example above to copy it to your clipboard!",
        "add_context": "Add reference material to help AI",
        "detected": "Detected",
        "content_type": "Content type",
        "modified": "Modified",
        "version": "Version",
        "size": "Size",
        "bytes": "bytes",
        "templates_available": "templates available",
        "search_files": "Search files",
        "select_file": "Select File",
        "select_template": "Select Template",
        "save_current": "Save Current Content",
        "connect": "Connect",
        "content_type_label": "Content Type",
        "context_optional": "Context (Optional)",
        "plugins_active": "Plugin(s) Active",
        "guide": "Guide",
        "no_images_library": "No images in library. Upload images to add them here.",
        "drag_drop": "Drag and drop file here",
        "browse_files": "Browse files",
        "syntax_inserted": "This syntax is inserted automatically when you click the plugin button above",
        "ai_engine": "AI Engine",
        "api_key": "API Key",
        "model": "Model",
        "paste_notes": "Paste notes here:",
        "file_name_placeholder": "Enter file name (without .md)",
        "save_button": "💾 Save",
        "choose_theme_help": "Choose a color theme for your content",
        "plugin_components_hint": "💡 Plugin components appear in 'Add Components' section above - just click to use!",
        "reload_plugins_help": "Reload plugins after creating new ones",
        "view_plugin_guide_help": "View plugin creation guide",
        "image_source_help": "Default source for [IMG: prompt] tags",
        "file_upload_limit": "Limit 200MB per file • PNG, JPG, JPEG, GIF",
        "preview_cached": "💡 Preview cached for performance ⚡ (cached)",
        "images_in_library": "{count} image(s) in library",
        "image_uploaded": "✅ Image uploaded!",
        "image_saved_to_library": "Image saved to library. You can reuse it from the Image Library above.",
        "preview_mode_help": "Mobile: WeChat/WeCom style | PC: Standard web style",
        "tab_visual": "Visual",
        "tab_wechat_code": "WeChat Code",
        "tab_standard_html": "Standard HTML",
        "choose_image_file": "Choose image file",
        "save_to_image_library": "Save to Image Library",
        "ai_engine_help": "Choose AI engine. 'None' disables AI features.",
        "plugin_component": "Plugin component",
        "ai_actions": "AI Actions",
        "generate_titles": "💡 Generate Titles",
        "expand_content": "📈 Expand Content",
        "smart_format": "✨ Smart Format",
        "suggest_components": "💡 Suggest Components",
        "polish_with_context": "🧠 Polish with Context",
        "polish_with_context_help": "Use AI to polish your content with context",
        "please_set_ai_engine": "Please set AI Engine at first.",
        "brainstorming_titles": "Brainstorming titles...",
        "detected_language": "Detected language: {lang} - Titles generated in {lang}",
        "expanding_content": "Expanding content...",
        "formatting_content": "Formatting content...",
        "analyzing_structure": "Analyzing structure...",
        "ai_processing": "AI Processing",
        "ai_working": "AI is working...",
        "polishing_content": "Polishing with context...",
        "ai_action_in_progress": "Another AI action is running.",
        "ai_action_debounced": "Please wait a moment before starting another AI action.",
        "ai_input_required": "Please add text before using this action.",
        "ai_action_failed": "AI request failed. Please retry.",
        "quick_toolbar": "Quick Access",
        "format_bold": "Bold",
        "format_italic": "Italic",
        "format_link": "Link",
        "format_code": "Code",
        "insert_hero": "Hero",
        "insert_card": "Card",
        "insert_2col": "2 Columns",
        "insert_3col": "3 Columns",
        "found_suggestions": "💡 Found {count} component suggestion(s)",
        "no_suggestions": "No component suggestions found.",
        "suggested_components": "💡 Suggested Components",
        "insert_component": "Insert component",
        "component_inserted": "✅ {name} inserted at {position}!",
        "clear_suggestions": "🗑️ Clear Suggestions",
        "configure_ai_engine": "⚠️ Please configure AI Engine in sidebar first.",
        "undo_help": "Undo last change",
        "redo_help": "Redo last undone change",
        "load_button": "📂 Load",
        "delete_button": "🗑️ Delete",
        "history_button": "📜 History",
        "file_info": "📊 {size} bytes • Modified: {date}",
        "copy_wechat_html": "📋 Copy WeChat HTML",
        "export_pdf": "📄 Export PDF",
        "export_word": "📝 Export Word",
        "download_html": "📥 Download HTML",
        "download_pdf": "📥 Download PDF",
        "download_word": "📥 Download Word",
    },
    "zh": {
        "app_title": "MarkPolish V1.0",
        "app_subtitle": "内容运营版",
        "simple_mode": "简洁模式",
        "dark_mode": "深色模式",
        "language": "语言",
        "ai_assistant": "AI 助手",
        "files_templates": "文件与模板",
        "add_components": "添加组件",
        "appearance": "外观设置",
        "plugins": "插件",
        "preview_mode": "预览模式",
        "editor": "编辑器",
        "preview": "预览",
        "mobile": "手机",
        "pc": "电脑",
        "words": "字",
        "chars": "字符",
        "min_read": "分钟阅读",
        "saved": "已保存",
        "saving": "保存中...",
        "save_error": "保存失败",
        "quick_save": "快速保存",
        "undo": "撤销",
        "redo": "重做",
        "export": "导出",
        "copy_wechat": "复制到微信",
        "standard_html": "标准 HTML",
        "download": "下载",
        "theme": "主题",
        "image_source": "图片来源",
        "files": "文件",
        "templates": "模板",
        "no_content": "暂无内容，请在编辑器中输入！",
        "layout": "布局",
        "content": "内容",
        "interactive": "交互",
        "media": "媒体",
        "comp_hero": "封面",
        "comp_2col": "两列",
        "comp_3col": "三列",
        "comp_steps": "步骤",
        "comp_timeline": "时间线",
        "comp_table": "表格",
        "comp_card": "卡片",
        "comp_reveal": "揭示",
        "comp_badge": "标签",
        "comp_button": "按钮",
        "comp_image": "AI 图片",
        "files": "文件",
        "templates": "模板",
        "save": "保存",
        "load": "加载",
        "delete": "删除",
        "version_history": "版本历史",
        "category": "分类",
        "all_templates": "全部模板",
        "use_template": "使用模板",
        "preview": "预览",
        "search_templates": "搜索模板...",
        "no_files": "暂无文件",
        "file_name": "文件名",
        "click_to_insert": "点击插入组件",
        "markdown_help": "Markdown 语法帮助",
        "image_assets": "图片与资源",
        "upload_image": "上传图片",
        "image_library": "图片库",
        "no_images": "暂无图片",
        "insert_image": "插入",
        "preview_mode": "预览模式",
        "mobile_preview": "手机",
        "pc_preview": "电脑",
        "reload_plugins": "刷新",
        "create_plugin": "创建插件",
        "installed_plugins": "已安装插件",
        "no_plugins": "暂无插件",
        "no_files_search": "没有匹配的文件",
        "no_files_yet": "暂无保存的文件。创建新文件或使用模板开始吧！",
        "no_templates_match": "没有匹配的模板",
        "click_to_insert_hint": "点击按钮插入组件。✅=微信兼容，🌐=仅 HTML",
        "view_full_guide": "查看完整指南",
        "syntax_tip": "提示：点击上方的语法示例可复制到剪贴板！",
        "add_context": "添加参考资料帮助 AI",
        "detected": "检测到",
        "content_type": "内容类型",
        "modified": "修改于",
        "version": "版本",
        "size": "大小",
        "bytes": "字节",
        "templates_available": "个模板可用",
        "search_files": "搜索文件",
        "select_file": "选择文件",
        "select_template": "选择模板",
        "save_current": "保存当前内容",
        "connect": "连接",
        "content_type_label": "内容类型",
        "context_optional": "上下文（可选）",
        "plugins_active": "个插件已激活",
        "guide": "指南",
        "no_images_library": "图片库为空，上传图片后将显示在此处。",
        "drag_drop": "拖放文件到这里",
        "browse_files": "浏览文件",
        "syntax_inserted": "点击上方插件按钮时，此语法将自动插入",
        "ai_engine": "AI 引擎",
        "api_key": "API 密钥",
        "model": "模型",
        "paste_notes": "在此粘贴笔记：",
        "file_name_placeholder": "输入文件名（不含 .md）",
        "save_button": "💾 保存",
        "choose_theme_help": "为内容选择颜色主题",
        "plugin_components_hint": "💡 插件组件出现在上方的「添加组件」部分 - 只需点击即可使用！",
        "reload_plugins_help": "创建新插件后重新加载",
        "view_plugin_guide_help": "查看插件创建指南",
        "image_source_help": "[IMG: prompt] 标签的默认来源",
        "file_upload_limit": "每个文件限制 200MB • PNG, JPG, JPEG, GIF",
        "preview_cached": "💡 预览已缓存以提升性能 ⚡（已缓存）",
        "images_in_library": "{count} 张图片在图库",
        "image_uploaded": "✅ 图片已上传！",
        "image_saved_to_library": "图片已保存到图库。可以在上方的图片库复用。",
        "preview_mode_help": "手机：微信/企业微信风格 | 电脑：标准网页风格",
        "tab_visual": "可视化",
        "tab_wechat_code": "微信代码",
        "tab_standard_html": "标准 HTML",
        "choose_image_file": "选择图片文件",
        "save_to_image_library": "保存到图片库",
        "ai_engine_help": "选择 AI 引擎。「无」将禁用 AI 功能。",
        "plugin_component": "插件组件",
        "ai_actions": "AI 操作",
        "generate_titles": "💡 生成标题",
        "expand_content": "📈 扩展内容",
        "smart_format": "✨ 智能格式化",
        "suggest_components": "💡 建议组件",
        "polish_with_context": "🧠 上下文润色",
        "polish_with_context_help": "使用 AI 根据上下文润色您的内容",
        "please_set_ai_engine": "请先设置 AI 引擎。",
        "brainstorming_titles": "正在生成标题...",
        "detected_language": "检测到语言：{lang} - 标题以 {lang} 生成",
        "expanding_content": "正在扩展内容...",
        "formatting_content": "正在格式化内容...",
        "analyzing_structure": "正在分析结构...",
        "ai_processing": "AI 处理中",
        "ai_working": "AI 正在处理...",
        "polishing_content": "正在结合上下文润色内容...",
        "ai_action_in_progress": "另一项 AI 操作正在运行。",
        "ai_action_debounced": "请稍等片刻再开始新的 AI 操作。",
        "ai_input_required": "请先输入内容再使用此操作。",
        "ai_action_failed": "AI 请求失败，请重试。",
        "quick_toolbar": "快速工具栏",
        "format_bold": "粗体",
        "format_italic": "斜体",
        "format_link": "链接",
        "format_code": "代码",
        "insert_hero": "英雄区",
        "insert_card": "卡片",
        "insert_2col": "2列",
        "insert_3col": "3列",
        "found_suggestions": "💡 找到 {count} 个组件建议",
        "no_suggestions": "未找到组件建议。",
        "suggested_components": "💡 建议的组件",
        "insert_component": "插入组件",
        "component_inserted": "✅ {name} 已插入到 {position}！",
        "clear_suggestions": "🗑️ 清除建议",
        "configure_ai_engine": "⚠️ 请先在侧边栏配置 AI 引擎。",
        "undo_help": "撤销上次更改",
        "redo_help": "重做上次撤销的更改",
        "load_button": "📂 加载",
        "delete_button": "🗑️ 删除",
        "history_button": "📜 历史",
        "file_info": "📊 {size} 字节 • 修改于：{date}",
        "copy_wechat_html": "📋 复制微信 HTML",
        "export_pdf": "📄 导出 PDF",
        "export_word": "📝 导出 Word",
        "download_html": "📥 下载 HTML",
        "download_pdf": "📥 下载 PDF",
        "download_word": "📥 下载 Word",
    }
}

def get_text(key):
    """Get translated text based on current language setting"""
    lang = st.session_state.get("ui_language", "en")
    return TRANSLATIONS.get(lang, TRANSLATIONS["en"]).get(key, key)

# Plugin name and description translations
PLUGIN_TRANSLATIONS = {
    "en": {
        "example_callout": "Example Callout",
        "example_callout_description": "Creates a callout box with different types (info, warning, error, success)",
    },
    "zh": {
        "example_callout": "示例标注",
        "example_callout_description": "创建不同类型的标注框（信息、警告、错误、成功）",
    }
}

def get_plugin_name(plugin_name):
    """Get translated plugin name"""
    lang = st.session_state.get("ui_language", "en")
    translations = PLUGIN_TRANSLATIONS.get(lang, PLUGIN_TRANSLATIONS["en"])
    key = plugin_name.lower().replace(" ", "_")
    return translations.get(key, plugin_name.replace('_', ' ').title())

def get_plugin_description(plugin_name, description):
    """Get translated plugin description"""
    if not description:
        return description
    lang = st.session_state.get("ui_language", "en")
    translations = PLUGIN_TRANSLATIONS.get(lang, PLUGIN_TRANSLATIONS["en"])
    key = f"{plugin_name.lower().replace(' ', '_')}_description"
    return translations.get(key, description)

# --- 0. SETUP ---
# Import configuration
try:
    from config import APP_TITLE, LAYOUT, INITIAL_SIDEBAR_STATE, setup_directories, TEMPLATES, SIDEBAR_HOVER_HTML
    # Setup page config using imported values
    st.set_page_config(page_title=APP_TITLE, layout=LAYOUT, initial_sidebar_state=INITIAL_SIDEBAR_STATE)
    setup_directories()
except ImportError:
    # Fallback if config module not available
    st.set_page_config(page_title="MarkPolish V1.0", layout="wide", initial_sidebar_state="collapsed")
    if not os.path.exists("projects"):
        os.makedirs("projects")
    if not os.path.exists("projects/images"):
        os.makedirs("projects/images")
    TEMPLATES = {}
    SIDEBAR_HOVER_HTML = ""  # Fallback empty string if config not available

# Import Helpers
try:
    from themes import STYLES
    from components import apply_components, INSERTION_TOOLS, COMPONENT_COMPATIBILITY
    from error_handler import ErrorHandler
    from performance import PerformanceOptimizer
    from plugin_system import get_plugin_registry, apply_plugins
except ImportError:
    STYLES = {}
    ErrorHandler = None
    PerformanceOptimizer = None
    COMPONENT_COMPATIBILITY = {}
    get_plugin_registry = None
    apply_plugins = None

try:
    from keyboard_listener import create_keyboard_listener
    HAS_KEYBOARD_LISTENER = True
except ImportError:
    HAS_KEYBOARD_LISTENER = False

# Import AI
try:
    from openai import OpenAI
    import httpx
    HAS_OPENAI = True
except ImportError:
    HAS_OPENAI = False

# Import new modules
try:
    from file_operations import (
        save_project, load_project, get_version_history, restore_version,
        get_version_storage_size, auto_save, load_auto_save, clear_auto_save,
        cleanup_old_autosave_files, push_to_undo_stack, undo_action, redo_action
    )
except ImportError:
    # Fallback if modules not available
    def save_project(name, content): return None, "Module not available"
    def load_project(filename): return None
    def get_version_history(project_name): return []
    def restore_version(project_name, version_index): return None, "Module not available"
    def get_version_storage_size(project_name): return 0
    def auto_save(content, project_name=None): return False, None
    def load_auto_save(project_name=None): return None, None
    def clear_auto_save(project_name=None): pass
    def cleanup_old_autosave_files(): pass
    def push_to_undo_stack(content, undo_stack, max_size=50): pass
    def undo_action(undo_stack, redo_stack, current_content): return None, undo_stack, redo_stack
    def redo_action(undo_stack, redo_stack, current_content): return None, undo_stack, redo_stack

try:
    from image_handling import (
        process_image, get_image_library, load_image_from_library, delete_image_from_library
    )
except ImportError:
    def process_image(file, save_to_library=True): return None, None
    def get_image_library(): return []
    def load_image_from_library(filename): return None
    def delete_image_from_library(filename): return False

try:
    from share_system import (
        generate_share_id, create_share_link, get_share_metadata, load_shared_project,
        get_share_link_url, list_project_shares, delete_share
    )
except ImportError:
    def generate_share_id(project_name): return None
    def create_share_link(project_name, permission="read", expires_days=30): return None, "Module not available"
    def get_share_metadata(share_id): return None
    def load_shared_project(share_id): return None, None
    def get_share_link_url(share_id): return ""
    def list_project_shares(project_name): return []
    def delete_share(share_id): return False

try:
    from ai_integration import check_connection, run_ai, detect_language
except ImportError:
    def check_connection(engine, url, key): return False, "Module not available"
    def run_ai(text, context, config, task_type="polish", content_type=None, available_plugins=None): return None, "Module not available"
    def detect_language(text): return "en"

try:
    from content_processing import (
        detect_content_type, get_preview_css, get_inline_styles, deep_inject_styles,
        parse_doc, clean_for_wechat, insert_component_at_position, get_stats
    )
except ImportError:
    def detect_content_type(text): return "article"
    def get_preview_css(theme, mode="Mobile"): return ""
    def get_inline_styles(theme): 
        t = theme if isinstance(theme, dict) else {'bg': '#fff', 'text': '#000', 'font': 'Arial', 'primary': '#4A90E2'}
        return {'wrapper': f"background-color: {t.get('bg', '#fff')}; padding: 20px; min-height: 100%; box-sizing: border-box;"}
    def deep_inject_styles(html_content, styles): return html_content
    def parse_doc(text, styles, img_provider="Pollinations (AI)", mode="wechat"): return text  # Return text as-is if module not available
    def clean_for_wechat(html): return html
    def insert_component_at_position(content, component_template, position): return content
    def get_stats(text): return 0, 0

try:
    from pdf_generator import generate_pdf, HAS_WEASYPRINT, HAS_PDFKIT, HAS_XHTML2PDF, HAS_REPORTLAB
except ImportError:
    def generate_pdf(html_content, theme, output_path=None, markdown_source=None, img_provider="Pollinations (AI)"): 
        return None, "PDF module not available"
    HAS_WEASYPRINT = False
    HAS_PDFKIT = False
    HAS_XHTML2PDF = False
    HAS_REPORTLAB = False

# Word Export
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

def generate_word(markdown_content, theme):
    """Generate Word document from markdown content - extracts text from components"""
    if not HAS_DOCX:
        return None, "Word export requires python-docx. Install with: pip install python-docx"
    
    try:
        from io import BytesIO
        import re as re_module
        
        doc = Document()
        
        # Get theme colors
        primary_hex = theme.get('primary', '#4A90E2').lstrip('#')
        text_hex = theme.get('text', '#333333').lstrip('#')
        
        def hex_to_rgb(hex_str):
            return tuple(int(hex_str[i:i+2], 16) for i in (0, 2, 4))
        
        primary_rgb = hex_to_rgb(primary_hex)
        text_rgb = hex_to_rgb(text_hex)
        
        # Pre-process: Clean ALL component syntax and extract only text
        def clean_components(text):
            """Strip ALL component syntax, HTML, and keep only plain text content"""
            # Remove ALL lines starting with ::: (component markers)
            text = re_module.sub(r'^:::.*$', '', text, flags=re_module.MULTILINE)
            
            # Remove [IMG: ...] placeholders (case insensitive)
            text = re_module.sub(r'\[IMG:.*?\]', '', text, flags=re_module.IGNORECASE)
            
            # Remove --split-- column separators
            text = text.replace('--split--', '')
            
            # Remove HTML tags but keep their text content
            text = re_module.sub(r'<[^>]+>', '', text)
            
            # Remove style attributes and other HTML-like syntax
            text = re_module.sub(r'style="[^"]*"', '', text)
            
            # Clean numbered list items in component format (1. text becomes just text in steps)
            # Keep standard markdown numbered lists
            
            # Clean up multiple blank lines
            text = re_module.sub(r'\n{3,}', '\n\n', text)
            
            # Remove lines that are just whitespace
            lines = [line for line in text.split('\n') if line.strip() or line == '']
            text = '\n'.join(lines)
            
            return text.strip()
        
        cleaned_content = clean_components(markdown_content)
        
        # Parse markdown line by line
        lines = cleaned_content.split('\n')
        current_list_items = []
        in_code_block = False
        code_content = []
        
        for line in lines:
            stripped = line.strip()
            
            # Handle code blocks
            if stripped.startswith('```'):
                if in_code_block:
                    # End code block
                    if code_content:
                        code_para = doc.add_paragraph()
                        code_run = code_para.add_run('\n'.join(code_content))
                        code_run.font.name = 'Courier New'
                        code_run.font.size = Pt(10)
                    code_content = []
                    in_code_block = False
                else:
                    in_code_block = True
                continue
            
            if in_code_block:
                code_content.append(line)
                continue
            
            # Skip empty lines
            if not stripped:
                # Flush list items if any
                if current_list_items:
                    for item in current_list_items:
                        doc.add_paragraph(item, style='List Bullet')
                    current_list_items = []
                continue
            
            # Headers
            if stripped.startswith('# '):
                p = doc.add_heading(stripped[2:], level=1)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(*primary_rgb)
            elif stripped.startswith('## '):
                p = doc.add_heading(stripped[3:], level=2)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(*primary_rgb)
            elif stripped.startswith('### '):
                p = doc.add_heading(stripped[4:], level=3)
            # List items
            elif stripped.startswith('- ') or stripped.startswith('* '):
                current_list_items.append(stripped[2:])
            elif re_module.match(r'^\d+\.\s', stripped):
                # Numbered list
                text = re_module.sub(r'^\d+\.\s', '', stripped)
                doc.add_paragraph(text, style='List Number')
            # Blockquote
            elif stripped.startswith('>'):
                quote_text = stripped[1:].strip()
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                run = p.add_run(quote_text)
                run.italic = True
            # Horizontal rule
            elif stripped in ['---', '***', '___']:
                doc.add_paragraph('_' * 50)
            # Regular paragraph
            else:
                # Flush list items first
                if current_list_items:
                    for item in current_list_items:
                        doc.add_paragraph(item, style='List Bullet')
                    current_list_items = []
                
                # Handle inline formatting
                p = doc.add_paragraph()
                
                # Parse inline bold, italic, links
                parts = re_module.split(r'(\*\*.*?\*\*|\*.*?\*|__.*?__|_.*?_|\[.*?\]\(.*?\))', stripped)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                        run.font.color.rgb = RGBColor(*primary_rgb)
                    elif part.startswith('__') and part.endswith('__'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                        run = p.add_run(part[1:-1])
                        run.italic = True
                    elif part.startswith('_') and part.endswith('_') and not part.startswith('__'):
                        run = p.add_run(part[1:-1])
                        run.italic = True
                    elif re_module.match(r'\[.*?\]\(.*?\)', part):
                        # Link - extract text and URL
                        link_match = re_module.match(r'\[(.*?)\]\((.*?)\)', part)
                        if link_match:
                            link_text = link_match.group(1)
                            run = p.add_run(link_text)
                            run.font.color.rgb = RGBColor(*primary_rgb)
                            run.underline = True
                    else:
                        p.add_run(part)
        
        # Flush remaining list items
        if current_list_items:
            for item in current_list_items:
                doc.add_paragraph(item, style='List Bullet')
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue(), "✅ Word document generated successfully"
    
    except Exception as e:
        return None, f"❌ Word generation error: {str(e)}"

# Import templates from config if not already imported or empty
if 'TEMPLATES' not in globals() or not TEMPLATES:
    try:
        from config import TEMPLATES
    except ImportError:
        TEMPLATES = {
            "Empty Draft": "",
            "⚡ Quick Update": "## Quick Update\n\nBrief announcement here.\n\n[Learn More](https://example.com)",
            "📌 Simple Notice": "## Important Notice\n\nPlease note the following:\n\n- Point 1\n- Point 2\n- Point 3\n\nThank you for your attention."
        }

# --- MAIN UI ---
# All helper functions have been moved to separate modules
# See: file_operations.py, image_handling.py, share_system.py, 
#      ai_integration.py, content_processing.py, pdf_generator.py

def main():
    # Check for share link in query parameters
    query_params = st.query_params
    
    # Read expander state from query params (set by JavaScript)
    # This allows JavaScript to update state
    if "expander_state" in query_params:
        # Check if we've already processed this state (prevent loops)
        last_processed = st.session_state.get("last_expander_state_hash", "")
        current_hash = str(query_params.get("expander_state", ""))
        
        if current_hash != last_processed:
            try:
                import json
                state_json = query_params["expander_state"]
                js_state = json.loads(state_json)
                # Initialize if not exists
                if "sidebar_expanded" not in st.session_state:
                    st.session_state.sidebar_expanded = {}
                # Update session state with JavaScript-detected state
                for key, value in js_state.items():
                    st.session_state.sidebar_expanded[key] = value
                # Mark as processed
                st.session_state.last_expander_state_hash = current_hash
            except Exception as e:
                pass  # Silently fail if JSON parsing fails
    if "share" in query_params:
        share_id = query_params["share"]
        shared_content, permission = load_shared_project(share_id)
        if shared_content is not None:
            st.session_state.content = shared_content
            st.session_state.editor_content = shared_content
            st.session_state.is_shared = True
            st.session_state.share_permission = permission
            st.session_state.share_id = share_id
            st.session_state.reset_editor = True
            # Reset preview cache to force re-render
            st.session_state.last_preview_content_hash = None
            if PerformanceOptimizer and st.session_state.get("performance_optimizer"):
                optimizer = st.session_state.performance_optimizer
                optimizer.preview_cache = {}
                optimizer.last_preview_hash = None
                optimizer.last_preview_time = 0
            # Clear query param after loading
            st.query_params.clear()
            st.rerun()
        else:
            st.error("⚠️ Share link is invalid or expired")
            st.query_params.clear()
    
    # Initialize session state variables
    if "auto_save_enabled" not in st.session_state:
        st.session_state.auto_save_enabled = True
    if "last_auto_save_time" not in st.session_state:
        st.session_state.last_auto_save_time = 0
    if "auto_save_status" not in st.session_state:
        st.session_state.auto_save_status = "saved"
    if "undo_stack" not in st.session_state:
        st.session_state.undo_stack = []
    if "redo_stack" not in st.session_state:
        st.session_state.redo_stack = []
    if "current_project_name" not in st.session_state:
        st.session_state.current_project_name = None
    if "keyboard_action" not in st.session_state:
        st.session_state.keyboard_action = None
    if "last_content_hash" not in st.session_state:
        st.session_state.last_content_hash = None
    if "performance_optimizer" not in st.session_state and PerformanceOptimizer:
        st.session_state.performance_optimizer = PerformanceOptimizer()
    if "preview_update_pending" not in st.session_state:
        st.session_state.preview_update_pending = False
    if "last_preview_content_hash" not in st.session_state:
        st.session_state.last_preview_content_hash = None
    if "is_shared" not in st.session_state:
        st.session_state.is_shared = False
    if "share_permission" not in st.session_state:
        st.session_state.share_permission = "read"
    
    # Cleanup old auto-save files on startup
    cleanup_old_autosave_files()
    
    # Check for auto-saved content on startup
    # Only show prompt if it hasn't been dismissed in this session
    if "restore_prompt_dismissed" not in st.session_state:
        st.session_state.restore_prompt_dismissed = False
    
    if not st.session_state.get("restore_prompt_dismissed", False):
        if "content" not in st.session_state or not st.session_state.get("content"):
            autosave_content, autosave_time = load_auto_save(st.session_state.current_project_name)
            if autosave_content and autosave_content.strip():
                st.session_state.show_restore_prompt = True
                st.session_state.restore_content = autosave_content
                st.session_state.restore_time = autosave_time
        else:
            st.session_state.show_restore_prompt = False
    else:
        st.session_state.show_restore_prompt = False
    
    # Re-inject sidebar hover code on every rerun to ensure it persists
    try:
        sidebar_html = SIDEBAR_HOVER_HTML
    except NameError:
        sidebar_html = ""  # Fallback if not imported
    if sidebar_html:
        st.components.v1.html(sidebar_html, height=0)
    
    # Inject JavaScript to track sidebar expander state
    expander_tracker_html = """
    <script>
    (function() {
        // Map of expander labels to state keys
        const expanderMap = {
            '🤖': 'ai_assistant',
            '📂': 'files_templates',
            '🧩': 'add_components',
            '🎨': 'appearance',
            '🔌': 'plugins',
            '🖼️': 'image_assets'
        };
        
        let lastState = {};
        let updateTimer = null;
        
        function getExpanderState() {
            const state = {};
            const expanders = document.querySelectorAll('[data-testid="stExpander"]');
            
            expanders.forEach(expander => {
                // Get the expander header
                const header = expander.querySelector('summary') || 
                             expander.querySelector('[data-testid="stExpanderToggleIcon"]')?.closest('div')?.parentElement;
                if (!header) return;
                
                const headerText = header.textContent || '';
                let expanderKey = null;
                
                // Match by emoji
                for (const [emoji, key] of Object.entries(expanderMap)) {
                    if (headerText.includes(emoji)) {
                        expanderKey = key;
                        break;
                    }
                }
                
                if (!expanderKey) return;
                
                // Check if expanded
                const toggleIcon = expander.querySelector('[data-testid="stExpanderToggleIcon"]');
                const isExpanded = expander.hasAttribute('open') || 
                                 (toggleIcon && toggleIcon.getAttribute('aria-expanded') === 'true');
                
                state[expanderKey] = isExpanded;
            });
            
            return state;
        }
        
        function updateState() {
            const currentState = getExpanderState();
            const stateChanged = JSON.stringify(currentState) !== JSON.stringify(lastState);
            
            if (stateChanged) {
                lastState = currentState;
                
                // Store in localStorage for persistence
                try {
                    localStorage.setItem('mp_sidebar_expanded', JSON.stringify(currentState));
                } catch(e) {}
                
                // Update URL with state (triggers Streamlit rerun)
                const stateJson = encodeURIComponent(JSON.stringify(currentState));
                const currentUrl = new URL(window.location);
                currentUrl.searchParams.set('expander_state', stateJson);
                
                // Use history.replaceState to avoid adding to history
                window.history.replaceState({}, '', currentUrl);
                
                // Trigger Streamlit rerun by dispatching a custom event
                window.dispatchEvent(new Event('popstate'));
            }
        }
        
        // Watch for expander clicks
        document.addEventListener('click', function(e) {
            if (e.target.closest('[data-testid="stExpander"]') || 
                e.target.closest('[data-testid="stExpanderToggleIcon"]')) {
                clearTimeout(updateTimer);
                updateTimer = setTimeout(updateState, 200);
            }
        }, true);
        
        // Watch for DOM changes
        const observer = new MutationObserver(function() {
            clearTimeout(updateTimer);
            updateTimer = setTimeout(updateState, 300);
        });
        
        function setupTracking() {
            const sidebar = document.querySelector('section[data-testid="stSidebar"]');
            if (sidebar) {
                observer.observe(sidebar, {
                    childList: true,
                    subtree: true,
                    attributes: true,
                    attributeFilter: ['open', 'aria-expanded']
                });
                
                // Load saved state from localStorage
                try {
                    const saved = localStorage.getItem('mp_sidebar_expanded');
                    if (saved) {
                        lastState = JSON.parse(saved);
                    }
                } catch(e) {}
                
                // Initial state check
                setTimeout(updateState, 500);
            } else {
                setTimeout(setupTracking, 200);
            }
        }
        
        setupTracking();
    })();
    </script>
    """
    st.components.v1.html(expander_tracker_html, height=0)
    
    with st.sidebar:
        st.title(get_text("app_title"))
        st.caption(get_text("app_subtitle"))
        
        # Settings Row: Simple Mode and Language Toggle
        settings_col1, settings_col2 = st.columns(2)
        with settings_col1:
            simple_mode = st.toggle(f"✨ {get_text('simple_mode')}", value=st.session_state.get("simple_mode", False), key="simple_mode")
        with settings_col2:
            # Language toggle: False = English, True = Chinese
            current_lang = st.session_state.get("ui_language", "en")
            is_chinese = current_lang == "zh"
            lang_toggle = st.toggle("🌐 中文", value=is_chinese, key="lang_toggle", help="Switch language / 切换语言")
            new_lang = "zh" if lang_toggle else "en"
            if new_lang != current_lang:
                st.session_state.ui_language = new_lang
                st.rerun()
        
        st.divider()
        
        # Initialize sidebar expander state memory
        if "sidebar_expanded" not in st.session_state:
            st.session_state.sidebar_expanded = {
                "ai_assistant": not simple_mode,
                "files_templates": False,
                "add_components": not simple_mode,
                "appearance": not simple_mode,
                "plugins": False,
                "image_assets": False
            }
        
        
        # 1. AI ASSISTANT (Most important - at the top)
        ai_expanded = st.session_state.sidebar_expanded.get("ai_assistant", not simple_mode)
        with st.expander(f"🤖 {get_text('ai_assistant')}", expanded=ai_expanded):
            engine_options = ["None", "Ollama (Local)", "OpenRouter"]
            prev_ai_cfg = st.session_state.get("ai_cfg", {"engine": "None", "key": "", "url": "", "model": ""})
            engine_default_idx = engine_options.index(prev_ai_cfg.get("engine", "None")) if prev_ai_cfg.get("engine", "None") in engine_options else 0

            engine = st.selectbox(
                get_text("ai_engine"), 
                engine_options,
                index=engine_default_idx,
                help=get_text("ai_engine_help"),
                key="ai_engine_select"
            )
            ai_cfg = {
                "engine": engine,
                "key": prev_ai_cfg.get("key", ""),
                "url": prev_ai_cfg.get("url", ""),
                "model": prev_ai_cfg.get("model", "")
            }
            
            if engine != "None":
                if engine == "OpenRouter":
                    ai_cfg["key"] = st.text_input(get_text("api_key"), value=ai_cfg.get("key", ""), type="password", help="Your OpenRouter API key", key="openrouter_api_key")
                    ai_cfg["url"] = "https://openrouter.ai/api/v1"
                    ai_cfg["model"] = st.text_input(get_text("model"), value=ai_cfg.get("model", "openai/gpt-4o-mini") or "openai/gpt-4o-mini", help="Model name (e.g., openai/gpt-4o-mini)", key="openrouter_model")
                    if st.button(f"🔌 {get_text('connect')}", use_container_width=True):
                        alive, msg = check_connection(engine, ai_cfg["url"], ai_cfg["key"])
                        if alive: 
                            st.success(msg)
                        else: 
                            st.error(msg)
                elif engine == "Ollama (Local)":
                    ai_cfg["url"] = st.text_input("Ollama URL", ai_cfg.get("url", "http://localhost:11434/v1") or "http://localhost:11434/v1", help="Local Ollama server URL", key="ollama_url")
                    ai_cfg["model"] = st.text_input(get_text("model"), value=ai_cfg.get("model", "llama3") or "llama3", help="Model name (e.g., llama3)", key="ollama_model")
                    if st.button(f"🔌 {get_text('connect')}", use_container_width=True):
                        alive, msg = check_connection(engine, ai_cfg["url"], None)
                        if alive: 
                            st.success(msg)
                        else: 
                            st.error(msg)
            
                # Content Type (simplified in simple mode)
                if not simple_mode:
                    current_content = st.session_state.get("content", "")
                    detected_type = detect_content_type(current_content) if current_content else None
                    
                    content_type_options = ["Auto-detect", "Product Launch", "Newsletter", "Tutorial", 
                                           "Marketing", "Internal", "Blog Post", "Announcement"]
                    default_idx = 0
                    if detected_type and detected_type in content_type_options:
                        default_idx = content_type_options.index(detected_type)
                    
                    selected_content_type = st.selectbox(
                        get_text("content_type_label"),
                        content_type_options,
                        index=default_idx,
                        help="Helps AI understand your content better"
                    )
                    
                    if selected_content_type == "Auto-detect":
                        ai_content_type = detected_type
                    else:
                        ai_content_type = selected_content_type if selected_content_type != "Auto-detect" else None
                    
                    if detected_type and selected_content_type == "Auto-detect":
                        st.caption(f"🔍 Detected: {detected_type}")
                    
                    # Context (advanced feature)
                    with st.expander(f"📝 {get_text('context_optional')}", expanded=False):
                        st.caption(get_text("add_context"))
                        context_text = st.text_area(get_text("paste_notes"), height=100, help="Meeting notes, requirements, etc.")
                else:
                    # Simple mode: auto-detect only
                    current_content = st.session_state.get("content", "")
                    ai_content_type = detect_content_type(current_content) if current_content else None
                    context_text = ""
                    if ai_content_type:
                        st.caption(f"🔍 Content type: {ai_content_type}")
            
            # Store AI config and content type in session state
            st.session_state.ai_cfg = ai_cfg
            st.session_state.ai_content_type = ai_content_type if engine != "None" else None
        
        # 2. FILES & TEMPLATES
        st.divider()
        files_expanded = st.session_state.sidebar_expanded.get("files_templates", False)
        with st.expander(f"📂 {get_text('files_templates')}", expanded=files_expanded):
            file_tab, template_tab = st.tabs([f"📁 {get_text('files')}", f"📄 {get_text('templates')}"])
            
            with file_tab:
                # File Management Section
                files = [f for f in os.listdir("projects") if f.endswith(".md") and not f.startswith(".")] if os.path.exists("projects") else []
                files = sorted(files, key=lambda x: os.path.getmtime(f"projects/{x}"), reverse=True)
                
                if files:
                    # Search box for files
                    file_search = st.text_input(f"🔍 {get_text('search_files')}", key="file_search", placeholder=get_text('search_files'))
                    filtered_files = [f for f in files if file_search.lower() in f.lower()] if file_search else files
                    
                    if filtered_files:
                        # File selection with better display
                        sel_file = st.selectbox(
                            get_text("select_file"),
                            ["New File"] + filtered_files,
                            format_func=lambda x: "📄 " + x.replace(".md", "") if x != "New File" else "➕ New File",
                            key="file_selector"
                        )
                        
                        # File actions
                        if sel_file != "New File":
                            file_col1, file_col2, file_col3 = st.columns([1, 1, 1])
                            with file_col1:
                                if st.button(get_text("load_button"), use_container_width=True, key="load_file_btn"):
                                    loaded_content, error = load_project(sel_file)
                                    if error:
                                        if ErrorHandler:
                                            ErrorHandler.show_error_with_details(error)
                                        else:
                                            st.error(error)
                                    elif loaded_content:
                                        st.session_state.content = loaded_content
                                        st.session_state.reset_editor = True
                                        project_name = sel_file.replace(".md", "")
                                        st.session_state.current_project_name = project_name
                                        clear_auto_save(project_name)
                                        if loaded_content:
                                            st.session_state.undo_stack = [loaded_content]
                                            st.session_state.redo_stack = []
                                        # Reset preview cache to force re-render after file load
                                        st.session_state.last_preview_content_hash = None
                                        if PerformanceOptimizer and st.session_state.get("performance_optimizer"):
                                            # Clear the preview cache
                                            optimizer = st.session_state.performance_optimizer
                                            optimizer.preview_cache = {}
                                            optimizer.last_preview_hash = None
                                            optimizer.last_preview_time = 0
                                        st.success(f"✅ Loaded {sel_file}")
                                        time.sleep(0.5)
                                        st.rerun()
                            
                            with file_col2:
                                if st.button(get_text("delete_button"), use_container_width=True, key="delete_file_btn"):
                                    try:
                                        os.remove(f"projects/{sel_file}")
                                        version_file = get_version_file_path(sel_file.replace(".md", ""))
                                        if os.path.exists(version_file):
                                            os.remove(version_file)
                                        st.success(f"✅ Deleted {sel_file}")
                                        time.sleep(0.5)
                                        st.rerun()
                                    except Exception as e:
                                        st.error(f"Failed to delete: {e}")
                            
                            with file_col3:
                                if st.button(get_text("history_button"), use_container_width=True, key="version_history_btn"):
                                    st.session_state.show_version_history = sel_file.replace(".md", "")
                                    st.rerun()
                            
                            # Show file info
                            file_path = f"projects/{sel_file}"
                            if os.path.exists(file_path):
                                file_size = os.path.getsize(file_path)
                                file_mtime = datetime.fromtimestamp(os.path.getmtime(file_path))
                                file_info_text = get_text("file_info").format(
                                    size=f"{file_size:,}",
                                    date=file_mtime.strftime('%Y-%m-%d %H:%M')
                                )
                                st.caption(file_info_text)
                    else:
                        st.info(get_text("no_files_search"))
                else:
                    st.info(get_text("no_files_yet"))
                
                st.divider()
                
                # Save Section
                st.subheader(f"💾 {get_text('save_current')}")
                save_name = st.text_input(get_text("file_name"), value=st.session_state.get("current_project_name", ""), key="save_name_input", placeholder=get_text("file_name_placeholder"))
                save_col1, save_col2 = st.columns([1, 1])
                with save_col1:
                    if st.button(get_text("save_button"), use_container_width=True, key="save_btn"):
                        if save_name:
                            if "content" in st.session_state:
                                result = save_project(save_name, st.session_state.content)
                                if result.startswith("✅"):
                                    st.success(result)
                                    if save_name:
                                        clear_auto_save(save_name)
                                    time.sleep(0.5)
                                    st.rerun()
                                else:
                                    if ErrorHandler:
                                        ErrorHandler.show_error_with_details(result)
                                    else:
                                        st.error(result)
                            else:
                                st.warning("No content to save")
                        else:
                            st.warning("Please enter a file name")
                
                with save_col2:
                    # Version History Settings
                    # Streamlit automatically updates session_state when checkbox value changes
                    st.checkbox(f"📜 {get_text('version_history')}", value=st.session_state.get("version_history_enabled", True), key="version_history_enabled")
            
            with template_tab:
                # Template Selection Section
                template_categories = {
                    "📝 Quick Start": ["Empty Draft", "⚡ Quick Update", "📌 Simple Notice"],
                    "📢 Product & Announcements": ["📢 Product Launch", "🎉 Feature Announcement", "📱 App Update"],
                    "📰 Content & News": ["📰 Weekly Newsletter", "📝 Blog Post", "📊 Industry Report"],
                    "🎯 Marketing & Promotions": ["🎯 Promotional Campaign", "💼 Case Study", "🎁 Special Offer"],
                    "📚 Educational": ["📚 Tutorial Guide", "🎓 How-To Article", "📖 FAQ Document"],
                    "🎪 Events": ["🎪 Event Announcement", "📅 Webinar Invitation", "🎊 Company Milestone"],
                    "👥 Internal & Team": ["📋 Meeting Summary", "📢 Internal Announcement", "🎯 Project Update"],
                    "💬 Customer & Support": ["💬 Customer Testimonial", "🆘 Support Guide", "🎁 Welcome Message"],
                    "🎨 Special Formats": ["📸 Photo Story", "🎬 Video Post", "📊 Data Report", "💡 Tip of the Day"]
                }
                
                # Show total template count
                total_templates = len(TEMPLATES) if TEMPLATES else 0
                st.caption(f"📚 {total_templates} {get_text('templates_available')}")
                
                # Category selector with counts
                category_options = [f"All Templates ({total_templates})"]
                for cat, templates in template_categories.items():
                    count = len([t for t in templates if t in TEMPLATES])
                    category_options.append(f"{cat} ({count})")
                
                selected_category_idx = st.selectbox(
                    get_text("category"),
                    category_options,
                    index=0,
                    key="template_category"
                )
                
                # Extract category name without count
                if "All Templates" in selected_category_idx:
                    selected_category = "All Templates"
                else:
                    selected_category = selected_category_idx.rsplit(" (", 1)[0]
                
                # Filter templates
                if selected_category == "All Templates":
                    available_templates = list(TEMPLATES.keys())
                else:
                    available_templates = template_categories[selected_category]
                
                # Template search
                template_search = st.text_input(f"🔍 {get_text('search_templates')}", key="template_search", placeholder=get_text('search_templates'))
                if template_search:
                    available_templates = [t for t in available_templates if template_search.lower() in t.lower()]
                
                if available_templates:
                    # Template selector
                    sel_template = st.selectbox(
                        get_text("select_template"),
                        available_templates,
                        index=0,
                        key="template_selector"
                    )
                    
                    # Template preview and use button
                    template_col1, template_col2 = st.columns([2, 1])
                    with template_col1:
                        if sel_template != "Empty Draft" and sel_template in TEMPLATES:
                            template_preview = TEMPLATES[sel_template]
                            preview_text = template_preview[:300] + "..." if len(template_preview) > 300 else template_preview
                            with st.expander(f"👁️ {get_text('preview')}", expanded=True):
                                st.code(preview_text, language="markdown")
                    
                    with template_col2:
                        if st.button(f"📄 {get_text('use_template')}", use_container_width=True, key="use_template_btn"):
                            if sel_template and sel_template in TEMPLATES:
                                template_content = TEMPLATES[sel_template]
                                st.session_state.content = template_content
                                st.session_state.reset_editor = True
                                # Reset preview cache to force re-render
                                st.session_state.last_preview_content_hash = None
                                if PerformanceOptimizer and st.session_state.get("performance_optimizer"):
                                    optimizer = st.session_state.performance_optimizer
                                    optimizer.preview_cache = {}
                                    optimizer.last_preview_hash = None
                                    optimizer.last_preview_time = 0

                                st.session_state.current_project_name = None
                                if template_content:
                                    st.session_state.undo_stack = [template_content]
                                    st.session_state.redo_stack = []
                                st.toast(f"✅ Loaded: {sel_template}")
                                st.rerun()
                            elif sel_template:
                                st.error(f"❌ Template '{sel_template}' not found. Available: {len(TEMPLATES)} templates")
                else:
                    st.info(get_text("no_templates_match"))
            
            # Version History Modal
            if st.session_state.get("show_version_history"):
                project_name = st.session_state.show_version_history
                st.divider()
                st.subheader(f"📜 {get_text('version_history')}: {project_name}")
                versions = get_version_history(project_name)
                
                if versions:
                    # Show version list
                    for idx, version in enumerate(reversed(versions)):
                        version_time = datetime.fromtimestamp(version.get("timestamp", 0))
                        version_size = version.get("size", 0)
                        
                        ver_col1, ver_col2, ver_col3 = st.columns([3, 1, 1])
                        with ver_col1:
                            st.write(f"**Version {len(versions) - idx}** - {version_time.strftime('%Y-%m-%d %H:%M:%S')}")
                            st.caption(f"Size: {version_size:,} bytes")
                        with ver_col2:
                            if st.button("📂 Restore", key=f"restore_v{idx}", use_container_width=True):
                                restored = restore_version(project_name, len(versions) - 1 - idx)
                                if restored:
                                    st.session_state.content = restored
                                    st.session_state.reset_editor = True
                                    # Reset preview cache to force re-render
                                    st.session_state.last_preview_content_hash = None
                                    if PerformanceOptimizer and st.session_state.get("performance_optimizer"):
                                        optimizer = st.session_state.performance_optimizer
                                        optimizer.preview_cache = {}
                                        optimizer.last_preview_hash = None
                                        optimizer.last_preview_time = 0

                                    st.success("✅ Version restored!")
                                    time.sleep(0.5)
                                    st.rerun()
                                else:
                                    st.error("Failed to restore version")
                        with ver_col3:
                            if st.button("👁️ View", key=f"view_v{idx}", use_container_width=True):
                                st.session_state.view_version_index = len(versions) - 1 - idx
                                st.session_state.view_version_project = project_name
                    
                    # Storage info
                    storage_size = get_version_storage_size(project_name)
                    st.caption(f"💾 Version history storage: {storage_size:,} bytes ({storage_size/1024:.1f} KB)")
                    
                    if st.button("🗑️ Clear History", key="clear_version_history"):
                        version_file = get_version_file_path(project_name)
                        if os.path.exists(version_file):
                            os.remove(version_file)
                        st.success("✅ Version history cleared")
                        time.sleep(0.5)
                        st.rerun()
                else:
                    st.info("No version history available for this file.")
                
                if st.button("❌ Close", key="close_version_history"):
                    st.session_state.show_version_history = None
                    st.rerun()
            
            # Migration Tool removed - users keep ## for headings, use ::: card for cards
            
            # Share link section
            if st.session_state.current_project_name:
                st.divider()
                st.subheader("🔗 Share Link")
                project_name = st.session_state.current_project_name
                
                # List existing shares
                existing_shares = list_project_shares(project_name)
                if existing_shares:
                    st.caption(f"{len(existing_shares)} active share link(s)")
                    for share in existing_shares[:3]:  # Show max 3
                        share_id = share.get("share_id")
                        permission = share.get("permission", "read")
                        expires_at = share.get("expires_at", 0)
                        expires_date = datetime.fromtimestamp(expires_at).strftime("%Y-%m-%d")
                        
                        col_share, col_del = st.columns([4, 1])
                        with col_share:
                            share_query = f"?share={share_id}"
                            st.code(share_query, language=None)
                            st.caption(f"{permission.upper()} • Expires: {expires_date}")
                        with col_del:
                            if st.button("🗑️", key=f"del_{share_id}", help="Delete share"):
                                if delete_share(share_id):
                                    st.success("Deleted")
                                    st.rerun()
                
                # Create new share link
                share_col1, share_col2 = st.columns(2)
                with share_col1:
                    share_permission = st.selectbox("Permission", ["read", "edit"], key="share_perm")
                with share_col2:
                    share_expires = st.selectbox("Expires", [7, 30, 90, 365], format_func=lambda x: f"{x} days", index=1, key="share_exp")
                
                if st.button("🔗 Create Share Link", use_container_width=True):
                    share_id, metadata = create_share_link(project_name, share_permission, share_expires)
                    if share_id:
                        st.success("✅ Share link created!")
                        share_query = f"?share={share_id}"
                        st.code(share_query, language=None)
                        st.info(f"Add `{share_query}` to your URL to share")
                        st.rerun()
                    else:
                        st.error("Failed to create share link")

        # 2. COMPONENTS (All component buttons in one expander)
        st.divider()
        components_expanded = st.session_state.sidebar_expanded.get("add_components", not simple_mode)
        with st.expander(f"🧩 {get_text('add_components')}", expanded=components_expanded):
            st.caption(get_text("click_to_insert"))
            
            # Group components by category (with translations)
            component_groups = {
                get_text("layout"): [
                    (get_text("comp_hero"), "::: hero\n# Title\nSubtitle\n:::"),
                    (get_text("comp_2col"), "::: col-2\nLeft\n--split--\nRight\n:::"),
                    (get_text("comp_3col"), "::: col-3\nOne\n--split--\nTwo\n--split--\nThree\n:::"),
                ],
                get_text("content"): [
                    (get_text("comp_steps"), "::: steps\n1. Step One\n2. Step Two\n:::"),
                    (get_text("comp_timeline"), "::: timeline\n2024 Start\n2025 Launch\n:::"),
                    (get_text("comp_table"), "::: table\nHeader 1 | Header 2 | Header 3\nRow 1 Col 1 | Row 1 Col 2 | Row 1 Col 3\n:::"),
                    (get_text("comp_card"), "::: card\n## Card Title\nCard content here.\n:::"),
                ],
                get_text("interactive"): [
                    (get_text("comp_reveal"), "::: reveal\nSecret Content\n--cover--\n👆\n:::"),
                    (get_text("comp_badge"), "[badge: NEW]"),
                    (get_text("comp_button"), "\n[Button Label](https://link.com)\n"),
                ],
                get_text("media"): [
                    (get_text("comp_image"), "[IMG: describe your image]"),
                ],
            }
            
            try:
                # Show grouped built-in components (now with translated names)
                layout_text = get_text("layout")
                content_text = get_text("content")
            
                for group_name, components in component_groups.items():
                    if not simple_mode or group_name in [layout_text, content_text]:
                        st.markdown(f"**{group_name}**")
                        
                        # Display in grid (3 columns)
                        cols_per_row = 3 if len(components) >= 3 else max(1, len(components))
                        for i in range(0, len(components), cols_per_row):
                            cols = st.columns(cols_per_row)
                            for j, col in enumerate(cols):
                                if i + j < len(components):
                                    comp_name, comp_syntax = components[i + j]
                                    def add_component(syntax=comp_syntax):
                                        if "content" not in st.session_state: 
                                            st.session_state.content = ""
                                        st.session_state.content += f"\n\n{syntax}\n"
                                        st.session_state.reset_editor = True
                                    col.button(comp_name, on_click=add_component, use_container_width=True, help=comp_name)
                        st.markdown("")  # Spacing between groups
                
                # Add plugin components section
                if get_plugin_registry:
                    try:
                        registry = get_plugin_registry()
                        plugins = registry.get_plugins_by_category()
                        
                        if plugins:
                            st.markdown("**🔌 Plugins**")
                            
                            # Show all plugins in a grid
                            plugin_list = list(plugins)
                            cols_per_row = 2
                            for i in range(0, len(plugin_list), cols_per_row):
                                cols = st.columns(cols_per_row)
                                for j, col in enumerate(cols):
                                    if i + j < len(plugin_list):
                                        plugin = plugin_list[i + j]
                                        if plugin.insertion_tool:
                                            plugin_display_name = get_plugin_name(plugin.name)
                                            compat_icon = "✅" if plugin.compatibility.get("wechat", False) else "🌐"
                                            
                                            def add_plugin_component(plugin_syntax=plugin.insertion_tool):
                                                if "content" not in st.session_state: 
                                                    st.session_state.content = ""
                                                st.session_state.content += f"\n\n{plugin_syntax}\n"
                                                st.session_state.reset_editor = True
                                        
                                            plugin_description = get_plugin_description(plugin.name, plugin.description)
                                            col.button(
                                                f"{plugin_display_name} {compat_icon}",
                                                on_click=add_plugin_component,
                                                use_container_width=True,
                                                help=f"{plugin_description or get_text('plugin_component')}"
                                            )
                    except Exception as e:
                        pass  # Silently fail for plugins
                            
            except Exception as e:
                st.error("Components not available")
            
            if not simple_mode:
                st.caption(f"💡 {get_text('click_to_insert_hint')}")

        # 3. APPEARANCE (Grouped settings)
        st.divider()
        appearance_expanded = st.session_state.sidebar_expanded.get("appearance", not simple_mode)
        with st.expander(f"🎨 {get_text('appearance')}", expanded=appearance_expanded):
            t_name = st.selectbox(
                get_text("theme"), 
                list(STYLES.keys()),
                help=get_text("choose_theme_help")
            )
        active_theme = STYLES[t_name]
        st.session_state.active_theme = active_theme 
        
        # Get theme and view (use session state if expander was collapsed)
        if st.session_state.get("active_theme"):
            active_theme = st.session_state.active_theme
        else:
            t_name = list(STYLES.keys())[0]
            active_theme = STYLES[t_name]
            st.session_state.active_theme = active_theme
        
        view = st.session_state.get("preview_view", "Mobile")
        
        # 5. PLUGINS (Plugin System) - Simplified UX
        if get_plugin_registry and not simple_mode:
            st.divider()
            plugins_expanded = st.session_state.sidebar_expanded.get("plugins", False)
            with st.expander(f"🔌 {get_text('plugins')}", expanded=plugins_expanded):
                try:
                    registry = get_plugin_registry()
                    plugins = registry.get_plugins_by_category()
                    
                    if plugins:
                        st.success(f"✅ **{len(plugins)} {get_text('plugins_active')}**")
                        st.caption(get_text("plugin_components_hint"))
                        
                        # Quick actions
                        col1, col2 = st.columns(2)
                        with col1:
                            if st.button(f"🔄 {get_text('reload_plugins')}", use_container_width=True, help=get_text("reload_plugins_help")):
                                try:
                                    from plugin_system import reload_plugin_registry
                                    reload_plugin_registry()
                                    st.success("Reloaded!")
                                    time.sleep(0.5)
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"Failed: {e}")
                        with col2:
                            if st.button(f"📚 {get_text('guide')}", use_container_width=True, help=get_text("view_plugin_guide_help")):
                                st.session_state.show_plugin_docs = True
                                st.rerun()
                        
                        st.divider()
                        st.markdown(f"**📦 {get_text('installed_plugins')}**")
                        
                        # Show plugins in a simple list
                        for plugin in plugins:
                            compat = "✅ WeChat" if plugin.compatibility.get("wechat", False) else "🌐 HTML only"
                            plugin_display_name = get_plugin_name(plugin.name)
                            plugin_description = get_plugin_description(plugin.name, plugin.description)
                            with st.expander(f"**{plugin_display_name}** - {compat}", expanded=False):
                                if plugin_description:
                                    st.caption(plugin_description)
                                if plugin.insertion_tool:
                                    st.code(plugin.insertion_tool, language="markdown")
                                    st.caption(f"💡 {get_text('syntax_inserted')}")
                    else:
                        # No plugins - show simple getting started
                        st.info("**🚀 No plugins yet**")
                        st.markdown("""
                        **Want to add custom components?**
                        
                        1. Go to `plugins/` folder
                        2. Copy `example_callout.py` 
                        3. Edit it to create your component
                        4. Click "🔄 Reload" above
                        5. Your component appears in "Add Components"!
                        """)
                        
                        if st.button(f"📚 {get_text('view_full_guide')}", use_container_width=True):
                            st.session_state.show_plugin_docs = True
                            st.rerun()
                    
                    # Migration tool removed - users keep ## for headings
                    
                    # Plugin Documentation (shown in Plugins section)
                    if st.session_state.get("show_plugin_docs", False):
                        st.divider()
                        is_zh = st.session_state.get("ui_language") == "zh"
                        st.subheader("📚 插件文档" if is_zh else "📚 Plugin Documentation")
                        
                        if is_zh:
                            doc_tabs = st.tabs(["概述", "创建插件", "示例", "微信兼容性"])
                        else:
                            doc_tabs = st.tabs(["Overview", "Creating Plugins", "Examples", "WeChat Compatibility"])
                        
                        with doc_tabs[0]:
                            if is_zh:
                                st.markdown("""
                                ## 什么是插件？
                                
                                插件允许你创建自定义 Markdown 组件，扩展 MarkPolish Studio 的功能。
                                
                                ### 优势
                                - ✨ 无需修改核心代码即可添加新组件
                                - 🔄 可与团队共享组件
                                - 🎨 根据需求自定义
                                - 🚀 动态扩展功能
                                
                                ### 插件分类
                                - **布局**：影响页面结构的组件
                                - **内容**：文字、标注、引用等
                                - **交互**：需要用户交互的组件（仅 HTML）
                                - **媒体**：图片、视频、相册
                                - **数据**：图表、表格、分析
                                """)
                            else:
                                st.markdown("""
                                ## What are Plugins?
                            
                            Plugins allow you to create custom markdown components that extend MarkPolish Studio's functionality.
                            
                            ### Benefits
                            - ✨ Add new components without modifying core code
                            - 🔄 Share components with your team
                            - 🎨 Customize for your specific needs
                            - 🚀 Extend functionality dynamically
                            
                            ### Plugin Categories
                            - **Layout**: Components that affect page structure
                            - **Content**: Text, callouts, quotes, etc.
                            - **Interactive**: Components with user interaction (HTML only)
                            - **Media**: Images, videos, galleries
                            - **Data**: Charts, tables, analytics
                            """)
                        
                        with doc_tabs[1]:
                            if is_zh:
                                st.markdown("""
                                ## 创建插件
                                
                                ### 第一步：复制示例
                                1. 进入 `plugins/` 文件夹
                                2. 复制 `example_callout.py` 作为模板
                                3. 重命名为你的插件名（如 `my_component.py`）
                                
                                ### 第二步：定义元数据
                                ```python
                                PLUGIN_METADATA = {
                                    "name": "my_component",
                                    "syntax": r'(?is):::\s*mycomponent\\n(.*?)\\n:::',
                                    "description": "组件功能描述",
                                    "category": "Content",
                                    "compatibility": {
                                        "wechat": True,  # 是否支持微信？
                                        "html": True     # 是否支持 HTML？
                                    },
                                    "insertion_tool": "::: mycomponent\\n内容\\n:::"
                                }
                                ```
                                
                                ### 第三步：实现渲染函数
                                ```python
                                def render(match, styles, mode="wechat"):
                                    content = match.group(1)
                                    # 在这里构建 HTML
                                    return f'<div>{content}</div>'
                                ```
                                
                                ### 第四步：测试
                                1. 保存插件文件
                                2. 点击插件区域的"🔄 刷新"
                                3. 你的组件将出现在组件列表中
                                """)
                            else:
                                st.markdown("""
                                ## Creating a Plugin
                            
                            ### Step 1: Copy the Example
                            1. Go to the `plugins/` folder
                            2. Copy `example_callout.py` as a template
                            3. Rename it to your plugin name (e.g., `my_component.py`)
                            
                            ### Step 2: Define Metadata
                            ```python
                            PLUGIN_METADATA = {
                                "name": "my_component",
                                "syntax": r'(?is):::\s*mycomponent\\n(.*?)\\n:::',
                                "description": "What your component does",
                                "category": "Content",
                                "compatibility": {
                                    "wechat": True,  # Works in WeChat?
                                    "html": True     # Works in HTML?
                                },
                                "insertion_tool": "::: mycomponent\\nYour content\\n:::"
                            }
                            ```
                            
                            ### Step 3: Implement Render Function
                            ```python
                            def render(match, styles, mode="wechat"):
                                content = match.group(1)
                                # Build your HTML here
                                return f'<div>{content}</div>'
                            ```
                            
                            ### Step 4: Test
                            1. Save your plugin file
                            2. Click "🔄 Reload Plugins" in the Plugins section
                            3. Your component will appear in the component list
                            """)
                        
                        with doc_tabs[2]:
                            if is_zh:
                                st.markdown("""
                                ## 示例插件
                            
                                ### 简单标注
                            ```python
                            PLUGIN_METADATA = {
                                "name": "callout",
                                "syntax": r':::\s*callout\\n(.*?)\\n:::',
                                "compatibility": {"wechat": True, "html": True}
                            }
                            
                            def render(match, styles, mode):
                                content = match.group(1)
                                return f'<div style="border-left: 4px solid {styles["primary"]}; padding: 10px;">{content}</div>'
                            ```
                                """)
                            else:
                                st.markdown("""
                                ## Example Plugins
                            
                                ### Simple Callout
                            ```python
                            PLUGIN_METADATA = {
                                    "name": "callout",
                                    "syntax": r':::\s*callout\\n(.*?)\\n:::',
                                    "compatibility": {"wechat": True, "html": True}
                            }
                            
                            def render(match, styles, mode):
                                    content = match.group(1)
                                    return f'<div style="border-left: 4px solid {styles["primary"]}; padding: 10px;">{content}</div>'
                            ```
                            """)
                        
                        with doc_tabs[3]:
                            if is_zh:
                                st.markdown("""
                                ## 微信兼容性指南
                                
                                ### ✅ 微信支持的功能
                                - 基础 HTML 标签（div, span, p, h1-h6, img, a）
                                - 内联 CSS（颜色、字体、基础布局）
                                - 静态图片（base64 或托管 URL）
                                - 简单布局（flexbox、基础定位）
                                
                                ### ❌ 微信不支持的功能
                                - JavaScript（无交互组件）
                                - 高级 CSS（动画、变换、滤镜）
                                - 外部资源（字体、脚本）
                                - iframe
                                - 复杂 SVG 动画
                                
                                ### 最佳实践
                                1. **在微信模式下测试** 后再标记为兼容
                                2. **使用内联样式** 而非 class
                                3. **避免 JavaScript** - 使用纯 CSS 方案
                                4. **保持简单** - 微信对 CSS 支持有限
                                5. **提供降级方案** - 组件不可用时显示提示
                                """)
                            else:
                                st.markdown("""
                                ## WeChat Compatibility Guide
                            
                            ### ✅ What Works in WeChat
                            - Basic HTML tags (div, span, p, h1-h6, img, a)
                            - Inline CSS (colors, fonts, basic layout)
                            - Static images (base64 or hosted URLs)
                            - Simple layouts (flexbox, basic positioning)
                            
                            ### ❌ What Doesn't Work
                            - JavaScript (no interactive components)
                            - Advanced CSS (animations, transforms, filters)
                            - External resources (fonts, scripts)
                            - iframes
                            - Complex SVG animations
                            
                            ### Best Practices
                            1. **Test in WeChat mode** before marking as compatible
                            2. **Use inline styles** instead of classes
                            3. **Avoid JavaScript** - use CSS-only solutions
                            4. **Keep it simple** - WeChat has limited CSS support
                            5. **Provide fallbacks** - show a message if component can't work
                                """)
                        
                        close_text = "❌ 关闭文档" if is_zh else "❌ Close Documentation"
                        if st.button(close_text, key="close_plugin_docs", use_container_width=True):
                            st.session_state.show_plugin_docs = False
                            st.rerun()
                        
                except Exception as e:
                    st.error(f"Plugin error: {e}")
                    if ErrorHandler:
                        ErrorHandler.log_error("plugin_ui", e)
        
        # 6. IMAGES & ASSETS (Grouped)
        st.divider()
        images_expanded = st.session_state.sidebar_expanded.get("image_assets", False)
        with st.expander(f"🖼️ {get_text('image_assets')}", expanded=images_expanded):
            img_provider = st.selectbox(
                get_text("image_source"), 
                ["Pollinations (AI)", "Picsum (Stock)", "Placeholder (Text)"],
                help=get_text("image_source_help")
            )
            
            # Image Library (nested in Images & Assets)
            st.markdown(f"**📚 {get_text('image_library')}**")
            library_images = get_image_library()
            if library_images:
                st.caption(get_text("images_in_library").format(count=len(library_images)))
                
                # Display images in grid
                cols_per_row = 3
                for i in range(0, len(library_images), cols_per_row):
                    cols = st.columns(cols_per_row)
                    for j, col in enumerate(cols):
                        if i + j < len(library_images):
                            img_info = library_images[i + j]
                            img_data = load_image_from_library(img_info["filename"])
                            
                            if img_data:
                                # Display thumbnail
                                col.image(img_data, use_container_width=True)
                                
                                # Image name and actions
                                display_name = img_info.get("original_name", img_info["filename"])[:20]
                                col.caption(display_name)
                                
                                # Insert and delete buttons
                                btn_col1, btn_col2 = col.columns(2)
                                with btn_col1:
                                    if st.button("➕", key=f"insert_{img_info['filename']}", help="Insert into editor"):
                                        if "content" not in st.session_state:
                                            st.session_state.content = ""
                                        # Use original name for reference
                                        ref_name = img_info.get("original_name", img_info["filename"])
                                        # Load into session state if not already there
                                        if "local_images" not in st.session_state:
                                            st.session_state.local_images = {}
                                        st.session_state.local_images[ref_name] = img_data
                                        st.session_state.content += f"\n[LOCAL: {ref_name}]\n"
                                        st.rerun()
                                with btn_col2:
                                    if st.button("🗑️", key=f"del_{img_info['filename']}", help="Delete"):
                                        if delete_image_from_library(img_info["filename"]):
                                            st.success("Deleted")
                                            st.rerun()
            else:
                st.info(get_text("no_images_library"))
        
            
            # Upload new image
            st.markdown(f"**📤 {get_text('upload_image')}**")
            uploaded_img = st.file_uploader(
                get_text("choose_image_file"), 
                type=['png', 'jpg', 'jpeg', 'gif'], 
                label_visibility="collapsed",
                help=get_text("file_upload_limit")
            )
            # Inject JavaScript to replace Streamlit's hardcoded file uploader text
            drag_drop_text = get_text("drag_drop").replace("'", "\\'").replace('"', '\\"').replace('\n', '\\n')
            browse_files_text = get_text("browse_files").replace("'", "\\'").replace('"', '\\"').replace('\n', '\\n')
            file_upload_limit_text = get_text("file_upload_limit").replace("'", "\\'").replace('"', '\\"').replace('\n', '\\n')
            # Custom visible Chinese labels (always rendered) — we hide Streamlit's built-ins
            st.markdown(
                f"""
                <div id="mp-uploader-labels" style="margin-top: -6px; margin-bottom: 10px;">
                    <div style="font-weight: 600; color: #1f1f1f; font-size: 14px;">{get_text("drag_drop")}</div>
                    <div style="color: #666; font-size: 12px; margin-top: 4px;">{get_text("file_upload_limit")}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )

            st.markdown(f"""
            <script>
            (function() {{
                const dragDropText = '{drag_drop_text}';
                const browseFilesText = '{browse_files_text}';
                const fileLimitText = '{file_upload_limit_text}';
                
                function replaceFileUploaderText() {{
                    // Method 1: Target specific "Browse files" button by testid
                    const browseButtons = document.querySelectorAll('[data-testid="stBaseButton-secondary"]');
                    browseButtons.forEach(btn => {{
                        if (btn.textContent && btn.textContent.trim() === 'Browse files') {{
                            btn.textContent = browseFilesText;
                            // Also replace in child text nodes
                            const walker = document.createTreeWalker(btn, NodeFilter.SHOW_TEXT, null, false);
                            let textNode;
                            while (textNode = walker.nextNode()) {{
                                if (textNode.textContent.includes('Browse files')) {{
                                    textNode.textContent = browseFilesText;
                                }}
                            }}
                        }}
                    }});
                    
                    // Method 2: Find all buttons and replace "Browse files" text
                    const allButtons = document.querySelectorAll('button');
                    allButtons.forEach(btn => {{
                        if (btn.textContent && btn.textContent.trim() === 'Browse files') {{
                            btn.textContent = browseFilesText;
                        }}
                    }});
                    
                    // Method 3: Find all file uploader areas
                    const uploadAreas = document.querySelectorAll('[data-testid="stFileUploader"]');
                    uploadAreas.forEach(uploadArea => {{
                        if (!uploadArea) return;
                    
                    // Force-set button and label texts (avoid relying on original English)
                    const buttons = uploadArea.querySelectorAll('button');
                    buttons.forEach(btn => {{
                        if (btn.textContent !== browseFilesText) {{
                            btn.textContent = browseFilesText;
                        }}
                    }});
                    
                    const labelNodes = uploadArea.querySelectorAll('span, p, div');
                    labelNodes.forEach(el => {{
                        const txt = (el.textContent || '').trim();
                        if (txt.startsWith('Drag and drop') || txt === dragDropText) {{
                            el.textContent = '';
                        }}
                        if (/^Limit 200MB per file/i.test(txt) || txt === fileLimitText) {{
                            el.textContent = '';
                        }}
                    }});
                        
                        // Find and replace all text nodes recursively
                        function replaceTextInNode(node) {{
                            if (node.nodeType === Node.TEXT_NODE) {{
                                if (node.textContent.includes('Drag and drop file here')) {{
                                node.textContent = '';
                                }}
                                if (node.textContent.includes('Browse files')) {{
                                node.textContent = browseFilesText;
                                }}
                                if (node.textContent.includes('Limit 200MB per file')) {{
                                node.textContent = '';
                                }}
                            }} else if (node.nodeType === Node.ELEMENT_NODE) {{
                                // Also check element textContent
                                if (node.textContent && node.textContent.trim()) {{
                                    if (node.textContent.includes('Drag and drop file here')) {{
                                        const walker = document.createTreeWalker(node, NodeFilter.SHOW_TEXT, null, false);
                                        let textNode;
                                        while (textNode = walker.nextNode()) {{
                                            if (textNode.textContent.includes('Drag and drop file here')) {{
                                            textNode.textContent = '';
                                            }}
                                        }}
                                    }}
                                    if (node.textContent.includes('Browse files')) {{
                                        const walker = document.createTreeWalker(node, NodeFilter.SHOW_TEXT, null, false);
                                        let textNode;
                                        while (textNode = walker.nextNode()) {{
                                            if (textNode.textContent.includes('Browse files')) {{
                                            textNode.textContent = browseFilesText;
                                            }}
                                        }}
                                    }}
                                    if (node.textContent.includes('Limit 200MB per file')) {{
                                        const walker = document.createTreeWalker(node, NodeFilter.SHOW_TEXT, null, false);
                                        let textNode;
                                        while (textNode = walker.nextNode()) {{
                                            if (textNode.textContent.includes('Limit 200MB per file')) {{
                                            textNode.textContent = '';
                                            }}
                                        }}
                                    }}
                                }}
                                
                                // Recursively process child nodes
                                Array.from(node.childNodes).forEach(child => replaceTextInNode(child));
                            }}
                        }}
                        
                        // Process the entire upload area
                        replaceTextInNode(uploadArea);
                    }});
                }}
                
                // Use MutationObserver to watch for DOM changes - observe entire body
                const observer = new MutationObserver(function(mutations) {{
                    let shouldRun = false;
                    mutations.forEach(mutation => {{
                        if (mutation.type === 'childList' || mutation.type === 'characterData') {{
                            shouldRun = true;
                        }}
                    }});
                    if (shouldRun) {{
                        setTimeout(replaceFileUploaderText, 10);
                    }}
                }});
                
                // Observe document body for all changes (more aggressive)
                observer.observe(document.body, {{ 
                    childList: true, 
                    subtree: true, 
                    characterData: true,
                    attributes: false
                }});
                
                // Also observe all upload areas individually
                document.querySelectorAll('[data-testid="stFileUploader"]').forEach(uploadArea => {{
                    observer.observe(uploadArea, {{ 
                        childList: true, 
                        subtree: true, 
                        characterData: true,
                        attributes: false
                    }});
                }});
                
                // Run immediately and repeatedly with more delays
                replaceFileUploaderText();
                [0, 50, 100, 200, 300, 500, 1000, 2000, 3000, 5000].forEach(delay => {{
                    setTimeout(replaceFileUploaderText, delay);
                }});
                
                // Also run after Streamlit reruns
                if (window.parent && window.parent.postMessage) {{
                    const originalPostMessage = window.parent.postMessage;
                    window.parent.postMessage = function(...args) {{
                        originalPostMessage.apply(this, args);
                        setTimeout(replaceFileUploaderText, 100);
                    }};
                }}
                
                // Run on every frame for a longer period after load
                let frameCount = 0;
                function frameCheck() {{
                    if (frameCount < 300) {{ // Check for ~5 seconds at 60fps
                        replaceFileUploaderText();
                        frameCount++;
                        requestAnimationFrame(frameCheck);
                    }}
                }}
                requestAnimationFrame(frameCheck);
                
                // Also add a continuous interval for 10 seconds
                let intervalCount = 0;
                const longInterval = setInterval(() => {{
                    replaceFileUploaderText();
                    intervalCount++;
                    if (intervalCount > 100) {{ // 10 seconds at 100ms intervals
                        clearInterval(longInterval);
                    }}
                }}, 100);
            }})();
            </script>
            """, unsafe_allow_html=True)
        if uploaded_img:
            save_to_lib = st.checkbox(get_text("save_to_image_library"), value=True, key="save_to_lib")
            shortcode, error = process_image(uploaded_img, save_to_library=save_to_lib)
            if error:
                if ErrorHandler:
                    ErrorHandler.show_error_with_details(error)
                else:
                    st.error(error)
            elif shortcode:
                st.success(get_text("image_uploaded"))
                st.code(shortcode, language="text")
                if save_to_lib:
                    st.info(get_text("image_saved_to_library"))
            else:
                st.error(shortcode if shortcode else "Upload failed")

        # 7. PREVIEW MODE (At the bottom)
        st.divider()
        view = st.radio(
            f"📱 {get_text('preview_mode')}", 
            [get_text('mobile_preview'), get_text('pc_preview')], 
            horizontal=True,
            help=get_text("preview_mode_help")
        )
        # Map back to English for internal use
        st.session_state.preview_view = "Mobile" if view == get_text('mobile_preview') else "PC"

    # Restore auto-save prompt
    if st.session_state.get("show_restore_prompt", False):
        restore_content = st.session_state.get("restore_content", "")
        restore_time = st.session_state.get("restore_time", 0)
        if restore_time:
            try:
                if hasattr(restore_time, "timestamp"):
                    ts_val = restore_time.timestamp()
                else:
                    ts_val = float(restore_time)
                time_str = datetime.fromtimestamp(ts_val).strftime("%Y-%m-%d %H:%M:%S")
            except Exception:
                time_str = "recently"
        else:
            time_str = "recently"
        
        with st.container():
            st.info(f"📋 Auto-saved content found from {time_str}. Restore?")
            col_restore, col_dismiss = st.columns(2)
            with col_restore:
                if st.button("✅ Restore", use_container_width=True):
                    st.session_state.content = restore_content
                    st.session_state.reset_editor = True  # Flag to reset editor on rerun
                    st.session_state.show_restore_prompt = False
                    st.rerun()
            with col_dismiss:
                if st.button("❌ Dismiss", use_container_width=True, key="dismiss_restore"):
                    st.session_state.show_restore_prompt = False
                    st.session_state.restore_prompt_dismissed = True  # Mark as dismissed to prevent re-showing
                    # Clear restore data to prevent it from showing again
                    if "restore_content" in st.session_state:
                        del st.session_state.restore_content
                    if "restore_time" in st.session_state:
                        del st.session_state.restore_time
                    st.rerun()

    col1, col2 = st.columns([1, 1])

    with col1:
        # Editor Column
        current_content = st.session_state.get("content", "")
        
        # Initialize editor_content if it doesn't exist or if content was updated externally
        if "editor_content" not in st.session_state:
            st.session_state.editor_content = current_content
        
        st.subheader(f"✒️ {get_text('editor')}")

        # Inline markdown syntax help
        with st.expander(f"📖 {get_text('markdown_help')}", expanded=False):
            help_col1, help_col2 = st.columns(2)
            
            with help_col1:
                st.markdown("**Basic Formatting**")
                st.code("""# H1 Header
## H2 Header
### H3 Header

**bold text**
*italic text*
~~strikethrough~~

`inline code`

[Link text](https://example.com)
![Image alt](image.jpg)""", language="markdown")
                
                st.markdown("**Lists**")
                st.code("""- Unordered item
- Another item
  - Nested item

1. Ordered item
2. Another item
   1. Nested item""", language="markdown")
            
            with help_col2:
                st.markdown("**Advanced**")
                st.code("""> Blockquote text
> Can span multiple lines

```python
# Code block
def hello():
    print("Hello")
```

| Column 1 | Column 2 |
|----------|----------|
| Cell 1   | Cell 2   |""", language="markdown")
                
                st.markdown("**Special Components**")
                st.code("""::: hero
# Title
Subtitle
:::

::: col-2
Left column
--split--
Right column
:::

[badge: NEW]
[Button Label](https://link.com)
[IMG: describe your image]""", language="markdown")
            
            st.caption(f"💡 {get_text('syntax_tip')}")
        
        # Markdown validation (optional, shown if enabled)
        if st.session_state.get("show_validation", False) and current_content:
            if ErrorHandler:
                is_valid, issues = ErrorHandler.validate_markdown_syntax(current_content)
                if not is_valid:
                    with st.expander("⚠️ Markdown Validation Issues", expanded=True):
                        for issue in issues:
                            st.warning(issue)
                        st.caption("💡 These issues may affect rendering. Fix them for best results.")

        # Initialize undo stack with current content if empty
        if not st.session_state.undo_stack and current_content:
            st.session_state.undo_stack = [current_content]
        
        # Sync editor_content from content if needed (before widget creation)
        # This handles cases where content was updated externally (undo/redo, load file, etc.)
        if st.session_state.get("reset_editor", False):
            # Delete editor_content to force it to sync from content
            if "editor_content" in st.session_state:
                del st.session_state.editor_content
            st.session_state.reset_editor = False
        
        # Initialize editor_content from content if it doesn't exist
        if "editor_content" not in st.session_state:
            st.session_state.editor_content = current_content
        
        # Show share mode indicator
        if st.session_state.get("is_shared", False):
            permission = st.session_state.get("share_permission", "read")
            if permission == "read":
                st.info("👁️ **Read-only mode** - This is a shared link. You can view but not edit.")
            else:
                st.info("✏️ **Edit mode** - This is a shared link. You can view and edit.")

        # Apply any pending snippet insert before rendering the textarea
        if st.session_state.get("pending_insert"):
            base_text = st.session_state.get("editor_content") or st.session_state.get("content", "")
            st.session_state.editor_content = base_text + st.session_state.pending_insert
            st.session_state.pending_insert = None

        # Text area outside form to allow content change detection
        # Disable if shared and read-only
        is_readonly = st.session_state.get("is_shared", False) and st.session_state.get("share_permission", "read") == "read"
        txt = st.text_area(
            "MD", 
            st.session_state.editor_content, 
            height=600, 
            label_visibility="collapsed",
            key="editor_content",
            disabled=is_readonly
        )
        
        # --- Editor Status Bar ---
        editor_text = st.session_state.get("editor_content", "")
        
        # Calculate stats
        word_count = len(editor_text.split()) if editor_text.strip() else 0
        char_count = len(editor_text)
        # Reading time: ~200 words per minute for Chinese, ~250 for English
        read_minutes = max(1, round(word_count / 200)) if word_count > 0 else 0
        
        # Auto-save status
        save_status = st.session_state.get("auto_save_status", "saved")
        if save_status == "saving":
            save_icon = f"💾 {get_text('saving')}"
        elif save_status == "error":
            save_icon = f"⚠️ {get_text('save_error')}"
        else:
            save_icon = f"✓ {get_text('saved')}"
        
        # Display status bar
        status_col1, status_col2, status_col3 = st.columns([2, 2, 1])
        with status_col1:
            st.caption(f"📝 {word_count:,} {get_text('words')} · {char_count:,} {get_text('chars')}")
        with status_col2:
            st.caption(f"⏱️ ~{read_minutes} {get_text('min_read')}")
        with status_col3:
            st.caption(save_icon)
        
        # Check for content changes (for auto-save and undo stack)
        if "editor_content" in st.session_state:
            editor_txt = st.session_state.editor_content
            # Check if content actually changed
            content_hash = hash(editor_txt)
            if content_hash != st.session_state.get("last_content_hash"):
                # Content changed - update session state
                st.session_state.content = editor_txt
                
                # Push to undo stack (only if content actually changed)
                if editor_txt:
                    st.session_state.undo_stack = push_to_undo_stack(
                        editor_txt, 
                        st.session_state.undo_stack
                    )
                    # Clear redo stack on new edit
                    st.session_state.redo_stack = []
                
                st.session_state.last_content_hash = content_hash
                # Trigger debounced auto-save
                st.session_state.pending_autosave = True
                st.session_state.autosave_timer = time.time()
                st.session_state.auto_save_status = "saving"
        
        # Process pending auto-save (debounced - 2 seconds after last change)
        if st.session_state.get("pending_autosave", False):
            # Initialize timer if not set
            if "autosave_timer" not in st.session_state or st.session_state.autosave_timer is None:
                st.session_state.autosave_timer = time.time()
            
            elapsed = time.time() - st.session_state.autosave_timer
            if elapsed >= 2:  # 2 second debounce
                content_to_save = st.session_state.get("editor_content", current_content)
                if content_to_save and st.session_state.auto_save_enabled:
                    success, result = auto_save(content_to_save, st.session_state.current_project_name)
                    if success:
                        st.session_state.last_auto_save_time = result
                        st.session_state.auto_save_status = "saved"
                    else:
                        st.session_state.auto_save_status = "error"
                st.session_state.pending_autosave = False
                st.session_state.autosave_timer = None
        
        # Check for periodic auto-save (every 30 seconds)
        current_time = time.time()
        time_since_last_save = current_time - st.session_state.last_auto_save_time
        if time_since_last_save > 30 and not st.session_state.get("pending_autosave", False):
            content_to_save = st.session_state.get("editor_content", current_content)
            if st.session_state.auto_save_enabled and content_to_save:
                success, result = auto_save(content_to_save, st.session_state.current_project_name)
                if success:
                    st.session_state.last_auto_save_time = result
                    st.session_state.auto_save_status = "saved"
                else:
                    st.session_state.auto_save_status = "error"
        
        # Handle keyboard actions
        keyboard_action = st.session_state.get("keyboard_action")
        if keyboard_action:
            content_for_action = st.session_state.get("editor_content", current_content)
            if keyboard_action == "save":
                if st.session_state.current_project_name:
                    result = save_project(st.session_state.current_project_name, content_for_action)
                    if result.startswith("✅"):
                        st.toast("Saved!")
                    else:
                        if ErrorHandler:
                            ErrorHandler.show_error_with_details(result)
                        else:
                            st.error(result)
                else:
                    success, result = auto_save(content_for_action, None)
                    if success:
                        st.session_state.last_auto_save_time = result
                        st.session_state.auto_save_status = "saved"
                        st.toast("Auto-saved!")
                st.session_state.keyboard_action = None
            elif keyboard_action == "undo":
                if st.session_state.undo_stack and len(st.session_state.undo_stack) > 1:
                    new_content = undo_action(
                        st.session_state.undo_stack,
                        st.session_state.redo_stack,
                        content_for_action
                    )
                    st.session_state.content = new_content
                    st.session_state.reset_editor = True  # Flag to reset editor on rerun
                    st.session_state.keyboard_action = None
                    st.rerun()
            elif keyboard_action == "redo":
                if st.session_state.redo_stack:
                    new_content = redo_action(
                        st.session_state.undo_stack,
                        st.session_state.redo_stack,
                        content_for_action
                    )
                    st.session_state.content = new_content
                    st.session_state.reset_editor = True  # Flag to reset editor on rerun
                    st.session_state.keyboard_action = None
                    st.rerun()
        
        # Keyboard shortcuts - use custom component if available
        if HAS_KEYBOARD_LISTENER:
            create_keyboard_listener()
        
        # Quick action buttons (keyboard shortcuts via buttons)
        kb_col1, kb_col2, kb_col3 = st.columns(3)
        with kb_col1:
            if st.button(f"💾 {get_text('quick_save')}", use_container_width=True, key="kb_save", help=get_text("save_current")):
                content_to_save = st.session_state.get("editor_content", current_content)
                if st.session_state.current_project_name:
                    result = save_project(st.session_state.current_project_name, content_to_save)
                    if result.startswith("✅"):
                        st.success(result)
                    else:
                        if ErrorHandler:
                            ErrorHandler.show_error_with_details(result)
                        else:
                            st.error(result)
                else:
                    success, result = auto_save(content_to_save, None)
                    if success:
                        st.session_state.last_auto_save_time = result
                        st.session_state.auto_save_status = "saved"
                        st.toast("Auto-saved!")
        with kb_col2:
            undo_disabled = not (st.session_state.undo_stack and len(st.session_state.undo_stack) > 1)
            if st.button(f"↶ {get_text('undo')}", use_container_width=True, key="kb_undo", disabled=undo_disabled, help=get_text("undo_help")):
                if not undo_disabled:
                    content_for_undo = st.session_state.get("editor_content", current_content)
                    new_content = undo_action(
                        st.session_state.undo_stack,
                        st.session_state.redo_stack,
                        content_for_undo
                    )
                    st.session_state.content = new_content
                    st.session_state.reset_editor = True  # Flag to reset editor on rerun
                    st.rerun()
        with kb_col3:
            redo_disabled = not st.session_state.redo_stack
            if st.button(f"↷ {get_text('redo')}", use_container_width=True, key="kb_redo", disabled=redo_disabled, help=get_text("redo_help")):
                if not redo_disabled:
                    content_for_redo = st.session_state.get("editor_content", current_content)
                    new_content = redo_action(
                        st.session_state.undo_stack,
                        st.session_state.redo_stack,
                        content_for_redo
                    )
                    st.session_state.content = new_content
                    st.session_state.reset_editor = True  # Flag to reset editor on rerun
                    st.rerun()

        # Real-time preview - no form needed!
        # Preview updates automatically as you type
        
        # Get current text from editor
        current_txt = st.session_state.get("editor_content", current_content)
        
        # Sync content to session state for preview
        if current_txt != st.session_state.get("content", ""):
            st.session_state.content = current_txt
        
        # Get AI config and content type from session state
        ai_cfg = st.session_state.get("ai_cfg", {"engine": "None", "key": "", "url": "", "model": ""})
        ai_content_type = st.session_state.get("ai_content_type", None)
        
        # Enhanced AI Actions
        if "ai_busy" not in st.session_state:
            st.session_state.ai_busy = False
        if "pending_ai_action" not in st.session_state:
            st.session_state.pending_ai_action = None
        if "last_ai_action_ts" not in st.session_state:
            st.session_state.last_ai_action_ts = 0.0

        pending_ai_action = st.session_state.get("pending_ai_action")
        if st.session_state.ai_busy and pending_ai_action:
            current_txt = st.session_state.get("editor_content", current_content)

            if pending_ai_action == "generate_titles":
                if not current_txt:
                    st.toast(get_text("ai_input_required"))
                    st.session_state.pending_ai_action = None
                    st.session_state.ai_busy = False
                    st.session_state.last_ai_action_ts = time.time()
                else:
                    try:
                        with st.spinner(get_text("brainstorming_titles")):
                            titles, stat = run_ai(current_txt, "", ai_cfg, task_type="titles", content_type=ai_content_type)
                        st.info(titles)
                        detected_lang = detect_language(current_txt) if current_txt else "English"
                        lang_text = get_text("detected_language").format(lang=detected_lang)
                        st.caption(f"🌐 {lang_text}")
                    except Exception:
                        st.toast(get_text("ai_action_failed"))
                    finally:
                        st.session_state.pending_ai_action = None
                        st.session_state.ai_busy = False

            elif pending_ai_action == "expand_content":
                if not current_txt:
                    st.toast(get_text("ai_input_required"))
                    st.session_state.pending_ai_action = None
                    st.session_state.ai_busy = False
                    st.session_state.last_ai_action_ts = time.time()
                else:
                    try:
                        st.session_state.undo_stack = push_to_undo_stack(
                            current_txt,
                            st.session_state.undo_stack
                        )
                        st.session_state.redo_stack = []

                        with st.spinner(get_text("expanding_content")):
                            res, msg = run_ai(current_txt, context_text, ai_cfg, task_type="expand", content_type=ai_content_type)
                            st.session_state.content = res
                            st.session_state.reset_editor = True  # Trigger editor reset
                            if res:
                                st.session_state.undo_stack = push_to_undo_stack(
                                    res,
                                    st.session_state.undo_stack
                                )
                            st.toast(f"✅ {msg} - Expanded content applied to editor!")
                            time.sleep(0.5)
                    except Exception:
                        st.toast(get_text("ai_action_failed"))
                    finally:
                        st.session_state.pending_ai_action = None
                        st.session_state.ai_busy = False
                        st.rerun()

            elif pending_ai_action == "smart_format":
                if not current_txt:
                    st.toast(get_text("ai_input_required"))
                    st.session_state.pending_ai_action = None
                    st.session_state.ai_busy = False
                    st.session_state.last_ai_action_ts = time.time()
                else:
                    try:
                        st.session_state.undo_stack = push_to_undo_stack(
                            current_txt,
                            st.session_state.undo_stack
                        )
                        st.session_state.redo_stack = []
                    
                        available_plugins = []
                        if get_plugin_registry:
                            try:
                                registry = get_plugin_registry()
                                available_plugins = registry.get_plugins_by_category()
                            except Exception:
                                pass
                        
                        with st.spinner(get_text("formatting_content")):
                            res, msg = run_ai(current_txt, context_text, ai_cfg, task_type="format", 
                                              content_type=ai_content_type, available_plugins=available_plugins)
                            st.session_state.content = res
                            st.session_state.reset_editor = True  # Trigger editor reset
                            if res:
                                st.session_state.undo_stack = push_to_undo_stack(
                                    res,
                                    st.session_state.undo_stack
                                )
                            st.toast(msg)
                            time.sleep(0.5)
                    except Exception:
                        st.toast(get_text("ai_action_failed"))
                    finally:
                        st.session_state.pending_ai_action = None
                        st.session_state.ai_busy = False
                        st.rerun()

            elif pending_ai_action == "suggest_components":
                if not current_txt:
                    st.toast(get_text("ai_input_required"))
                    st.session_state.pending_ai_action = None
                    st.session_state.ai_busy = False
                    st.session_state.last_ai_action_ts = time.time()
                else:
                    try:
                        available_plugins = []
                        if get_plugin_registry:
                            try:
                                registry = get_plugin_registry()
                                available_plugins = registry.get_plugins_by_category()
                            except Exception:
                                pass
                        
                        with st.spinner(get_text("analyzing_structure")):
                            suggestions_raw, stat = run_ai(current_txt, "", ai_cfg, task_type="suggest_components", 
                                                          content_type=ai_content_type, available_plugins=available_plugins)
                        
                        import json
                        suggestions_list = []
                        try:
                            json_str = suggestions_raw.strip()
                            if json_str.startswith("```"):
                                json_str = json_str.split("```")[1]
                                if json_str.startswith("json"):
                                    json_str = json_str[4:]
                            json_str = json_str.strip()
                            suggestions_list = json.loads(json_str)
                        except Exception:
                            suggestions_lower = suggestions_raw.lower()
                            component_names = ["hero", "col-2", "col-3", "steps", "timeline", "card"]
                            if get_plugin_registry:
                                try:
                                    registry = get_plugin_registry()
                                    plugins = registry.get_plugins_by_category()
                                    for plugin in plugins:
                                        component_names.append(plugin.name)
                                except Exception:
                                    pass
                            for comp in component_names:
                                if comp in suggestions_lower:
                                    suggestions_list.append({"component": comp, "position": "end"})
                        
                        st.session_state.component_suggestions = suggestions_list
                        st.session_state.component_suggestions_content = current_txt  # Store content snapshot

                        if suggestions_list:
                            st.success(get_text("found_suggestions").format(count=len(suggestions_list)))
                        else:
                            st.info(get_text("no_suggestions"))
                    except Exception:
                        st.toast(get_text("ai_action_failed"))
                    finally:
                        st.session_state.pending_ai_action = None
                        st.session_state.ai_busy = False

            elif pending_ai_action == "polish_with_context":
                if not current_txt:
                    st.toast(get_text("ai_input_required"))
                    st.session_state.pending_ai_action = None
                    st.session_state.ai_busy = False
                    st.session_state.last_ai_action_ts = time.time()
                else:
                    try:
                        st.session_state.undo_stack = push_to_undo_stack(
                            current_txt,
                            st.session_state.undo_stack
                        )
                        st.session_state.redo_stack = []

                        with st.spinner(get_text("polishing_content")):
                            res, msg = run_ai(current_txt, context_text, ai_cfg, task_type="polish", content_type=ai_content_type)
                            st.session_state.content = res
                            st.session_state.reset_editor = True
                            if res:
                                st.session_state.undo_stack = push_to_undo_stack(
                                    res,
                                    st.session_state.undo_stack
                                )
                            st.toast(f"✅ {msg}")
                            time.sleep(0.5)
                    except Exception:
                        st.toast(get_text("ai_action_failed"))
                    finally:
                        st.session_state.pending_ai_action = None
                        st.session_state.ai_busy = False
                        st.rerun()

        ai_busy = st.session_state.ai_busy
        st.subheader(f"🤖 {get_text('ai_actions')}")

        def trigger_ai_action(action_name: str, require_text: bool = False):
            now = time.time()
            if st.session_state.ai_busy:
                st.toast(get_text("ai_action_in_progress"))
                return
            if now - st.session_state.get("last_ai_action_ts", 0) < 0.4:
                st.toast(get_text("ai_action_debounced"))
                return
            if ai_cfg['engine'] == "None":
                st.toast(get_text("please_set_ai_engine"))
                return
            if require_text and (not current_txt or not current_txt.strip()):
                st.toast(get_text("ai_input_required"))
                return
            st.session_state.pending_ai_action = action_name
            st.session_state.ai_busy = True
            st.session_state.last_ai_action_ts = now
            st.rerun()

        ai_col1, ai_col2 = st.columns(2)
        
        with ai_col1:
            if st.button(get_text("generate_titles"), use_container_width=True, disabled=ai_busy):
                trigger_ai_action("generate_titles", require_text=True)
            
            if st.button(get_text("expand_content"), use_container_width=True, disabled=ai_busy):
                trigger_ai_action("expand_content", require_text=True)
        
        with ai_col2:
            if st.button(get_text("smart_format"), use_container_width=True, disabled=ai_busy):
                trigger_ai_action("smart_format", require_text=True)
            
            if st.button(get_text("suggest_components"), use_container_width=True, disabled=ai_busy):
                trigger_ai_action("suggest_components", require_text=True)
            
            # Display stored suggestions with insert buttons (persists across reruns)
            if "component_suggestions" in st.session_state and st.session_state.component_suggestions:
                st.divider()
                st.subheader(get_text("suggested_components"))
                
                # Build component map with built-in components
                component_map = {
                    "hero": "::: hero\n# Title\nSubtitle\n:::",
                    "col-2": "::: col-2\nLeft content\n--split--\nRight content\n:::",
                    "col-3": "::: col-3\nOne\n--split--\nTwo\n--split--\nThree\n:::",
                    "steps": "::: steps\n1. Step One\n2. Step Two\n:::",
                    "timeline": "::: timeline\n2024 Event\n2025 Event\n:::",
                    "card": "::: card\n## Card Title\nCard content here.\n:::"
                }
                
                # Add plugin components to the map
                if get_plugin_registry:
                    try:
                        registry = get_plugin_registry()
                        plugins = registry.get_plugins_by_category()
                        for plugin in plugins:
                            if plugin.insertion_tool:
                                component_map[plugin.name] = plugin.insertion_tool
                    except:
                        pass
                
                # Show buttons for each suggestion
                for idx, suggestion in enumerate(st.session_state.component_suggestions):
                    comp_name = suggestion.get("component", "")
                    position = suggestion.get("position", "end")
                    
                    if comp_name in component_map:
                        comp_template = component_map[comp_name]
                        comp_display_name = comp_name.replace("-", " ").title()
                        
                        sug_col1, sug_col2 = st.columns([3, 1])
                        with sug_col1:
                            st.write(f"**{comp_display_name}** - *{position}*")
                        with sug_col2:
                            if st.button("➕", key=f"insert_comp_{idx}", use_container_width=True, help=get_text("insert_component")):
                                # Get current content
                                current_content = st.session_state.get("content", "")
                                
                                # Insert at the suggested position
                                new_content = insert_component_at_position(
                                    current_content, comp_template, position
                                )
                                
                                st.session_state.content = new_content
                                st.session_state.reset_editor = True
                                
                                # Remove this suggestion from the list
                                st.session_state.component_suggestions.pop(idx)
                                
                                inserted_msg = get_text("component_inserted").format(name=comp_display_name, position=position)
                                st.toast(inserted_msg)
                                time.sleep(0.3)
                                st.rerun()
                
                # Clear all suggestions button
                if st.button(get_text("clear_suggestions"), use_container_width=True):
                    st.session_state.component_suggestions = []
                    st.rerun()

        # AI Polish with Context button
        st.divider()
        if st.button(get_text("polish_with_context"), use_container_width=True, help=get_text("polish_with_context_help"), disabled=ai_busy):
            trigger_ai_action("polish_with_context", require_text=True)

    with col2:
        st.subheader(f"👁️ {view} Preview")
        
        # Get content from editor (prefer editor_content as it's the most up-to-date)
        content_to_render = st.session_state.get("editor_content") or st.session_state.get("content", "")
        
        # Initialize wechat_final to ensure it's always defined
        wechat_final = None
        
        # Performance optimization: Debounced preview rendering
        should_render_preview = True
        cached_preview = None
        
        if PerformanceOptimizer and st.session_state.get("performance_optimizer"):
            optimizer = st.session_state.performance_optimizer
            
            # Check if we should update preview (debouncing)
            # Include theme in hash so theme changes trigger re-render
            theme_key = str(active_theme.get('bg', '')) + str(active_theme.get('primary', ''))
            content_hash = hash(content_to_render + theme_key)
            if content_hash == st.session_state.get("last_preview_content_hash"):
                # Content AND theme haven't changed, use cached preview if available
                cached_preview = optimizer.get_cached_preview(content_to_render + theme_key)
                if cached_preview:
                    should_render_preview = False
            else:
                # Content or theme changed - check debounce
                if optimizer.should_update_preview(content_to_render):
                    st.session_state.last_preview_content_hash = content_hash
                    should_render_preview = True
                else:
                    # Use cached preview while debouncing
                    cached_preview = optimizer.get_cached_preview(content_to_render + theme_key)
                    if cached_preview:
                        should_render_preview = False
        
        # Render preview if there's content
        has_content = content_to_render and content_to_render.strip()
        if has_content:
            # Try to use cached preview first (if available and should_render_preview is False)
            if not should_render_preview and cached_preview:
                wechat_final = cached_preview
            else:
                # Render fresh preview
                try:
                    inline_styles = get_inline_styles(active_theme)
                    parsed_md = parse_doc(content_to_render, inline_styles, img_provider=img_provider, mode="wechat")
            
                    # Ensure parsed_md is not empty - if parse_doc returns empty, use original content
                    if not parsed_md or not parsed_md.strip():
                        parsed_md = content_to_render
            
                    # parsed_md contains HTML from plugins/components mixed with markdown
                    # markdown.markdown() escapes HTML, so we need to preserve HTML blocks
                    import re as re_module
            
                    # Strategy: Extract HTML blocks, process markdown, then merge back
                    # Use HTML comments as placeholders (markdown preserves them)
                    html_blocks = []
            
                    # Find all complete HTML elements
                    def extract_html(match):
                        html_block = match.group(0)
                        # Use HTML comment as placeholder (markdown preserves comments)
                        placeholder = f"<!--MPHTML{len(html_blocks)}-->"
                        html_blocks.append(html_block)
                        return placeholder
            
                    # Match complete HTML elements (section, div, span) with proper closing tags
                    html_pattern = r'<(section|div|span)[^>]*>.*?</\1>'
                    text_for_markdown = re_module.sub(html_pattern, extract_html, parsed_md, flags=re_module.DOTALL | re_module.IGNORECASE)
            
                    # Process markdown (HTML comments are preserved)
                    raw_html = markdown.markdown(text_for_markdown, extensions=['nl2br', 'extra'])
            
                    # Restore HTML blocks - replace placeholders with actual HTML
                    wechat_html_inner = raw_html
                    for i, html_block in enumerate(html_blocks):
                        placeholder = f"<!--MPHTML{i}-->"
                        wechat_html_inner = wechat_html_inner.replace(placeholder, html_block)
            
                    wechat_html_inner = deep_inject_styles(wechat_html_inner, inline_styles)
                    # Wrapper div: use theme background but NO padding (canvas already has padding: 20px)
                    wrapper_bg = active_theme.get('bg', '#fff')
                    wrapper_style = f"background-color: {wrapper_bg}; padding: 0; margin: 0; box-sizing: border-box;"
                    wechat_final = f'<div style="{wrapper_style}">{wechat_html_inner}</div>'
            
                    # Cache the preview (include theme in cache key)
                    if PerformanceOptimizer and st.session_state.get("performance_optimizer"):
                        cache_key = content_to_render + str(active_theme.get('bg', '')) + str(active_theme.get('primary', ''))
                        st.session_state.performance_optimizer.cache_preview(cache_key, wechat_final)
                except Exception as e:
                    # If rendering fails, try to show at least the raw markdown
                    try:
                        # Fallback: just render the markdown directly
                        raw_html_fallback = markdown.markdown(content_to_render, extensions=['nl2br', 'extra'])
                        inline_styles = get_inline_styles(active_theme)
                        # Don't add wrapper background - let the canvas handle it for proper theme matching
                        # Wrapper div: use theme background but NO padding (canvas already has padding: 20px)
                        wrapper_bg = active_theme.get('bg', '#fff')
                        wrapper_style = f"background-color: {wrapper_bg}; padding: 0; margin: 0; box-sizing: border-box;"
                        wechat_final = f'<div style="{wrapper_style}">{raw_html_fallback}</div>'
                    except:
                        # Last resort: show error message
                        wechat_final = f'<div style="padding: 40px; text-align: center; color: #999;"><p>Preview error. Content: {len(content_to_render)} chars</p></div>'
        else:
            # No content - show placeholder
            wechat_final = None
        
        # Ensure parsed_md is always defined for standard HTML generation
        if 'parsed_md' not in locals():
            if content_to_render and content_to_render.strip():
                inline_styles = get_inline_styles(active_theme)
                parsed_md = parse_doc(content_to_render, inline_styles, img_provider=img_provider, mode="wechat")
            else:
                parsed_md = ""
        
        t = active_theme
        # For standard HTML, also need to preserve HTML blocks
        import re as re_module_std
        html_blocks_std = []
        def extract_html_std(match):
            html_block = match.group(0)
            placeholder = f"<!--MPHTML{len(html_blocks_std)}-->"
            html_blocks_std.append(html_block)
            return placeholder
        parsed_md_for_std = re_module_std.sub(r'<(section|div|span)[^>]*>.*?</\1>', extract_html_std, parsed_md, flags=re_module_std.DOTALL | re_module_std.IGNORECASE)
        standard_html_content = markdown.markdown(parsed_md_for_std, extensions=['nl2br', 'extra'])
        for i, html_block in enumerate(html_blocks_std):
            standard_html_content = standard_html_content.replace(f"<!--MPHTML{i}-->", html_block)
        standard_full = f"""<!DOCTYPE html><html><head><style>body{{font-family:{t['font']};padding:20px;max-width:800px;margin:0 auto;line-height:1.7;color:{t['text']};background:{t['bg']};}} img{{max-width:100%;height:auto;}} a{{color:{t['primary']};}}</style></head><body>{standard_html_content}</body></html>"""
        
        # Performance indicator (subtle, only when using cache)
        preview_status = ""
        if not should_render_preview and cached_preview:
            preview_status = " ⚡ (cached)"
        
        t1, t2, t3 = st.tabs([get_text("tab_visual"), get_text("tab_wechat_code"), get_text("tab_standard_html")])
        with t1:
            import random
            k = random.randint(0,10000)
            if preview_status:
                st.caption(get_text("preview_cached"))
            # Ensure wechat_final has content, show placeholder if empty
            if not wechat_final:
                # No preview generated - check if we have content to render
                if content_to_render and content_to_render.strip():
                    # Content exists but preview failed - force render one more time
                    try:
                        inline_styles = get_inline_styles(active_theme)
                        parsed_md = parse_doc(content_to_render, inline_styles, img_provider=img_provider, mode="wechat")
                        
                        # Ensure parsed_md is not empty - if parse_doc returns empty, use original content
                        if not parsed_md or not parsed_md.strip():
                            parsed_md = content_to_render
                        
                        import re as re_module_final
                        html_blocks_final = []
                        def extract_html_final(match):
                            html_block = match.group(0)
                            placeholder = f"<!--MPHTML{len(html_blocks_final)}-->"
                            html_blocks_final.append(html_block)
                            return placeholder
                        text_for_markdown_final = re_module_final.sub(r'<(section|div|span)[^>]*>.*?</\1>', extract_html_final, parsed_md, flags=re_module_final.DOTALL | re_module_final.IGNORECASE)
                        raw_html_final = markdown.markdown(text_for_markdown_final, extensions=['nl2br', 'extra'])
                        wechat_html_inner_final = raw_html_final
                        for i, html_block in enumerate(html_blocks_final):
                            placeholder = f"<!--MPHTML{i}-->"
                            wechat_html_inner_final = wechat_html_inner_final.replace(placeholder, html_block)
                        wechat_html_inner_final = deep_inject_styles(wechat_html_inner_final, inline_styles)
                        wrapper_bg = active_theme.get('bg', '#fff')
                        wrapper_style = f"background-color: {wrapper_bg}; padding: 0; margin: 0; box-sizing: border-box;"
                        wechat_final = f'<div style="{wrapper_style}">{wechat_html_inner_final}</div>'
                    except Exception as e:
                        # If rendering fails, try simple markdown render
                        try:
                            raw_html_simple = markdown.markdown(content_to_render, extensions=['nl2br', 'extra'])
                            inline_styles = get_inline_styles(active_theme)
                            wrapper_bg = active_theme.get('bg', '#fff')
                            wrapper_style = f"background-color: {wrapper_bg}; padding: 0; margin: 0; box-sizing: border-box;"
                            wechat_final = f'<div style="{wrapper_style}">{raw_html_simple}</div>'
                        except:
                            # Last resort: show placeholder
                            no_content_text = get_text("no_content")
                            wechat_final = f'<div style="padding: 40px; text-align: center; color: #999;"><p>{no_content_text}</p></div>'
                else:
                    # No content - show placeholder
                    no_content_text = get_text("no_content")
                    wechat_final = f'<div style="padding: 40px; text-align: center; color: #999;"><p>{no_content_text}</p></div>'
            elif wechat_final.strip() == "":
                # Empty string - show placeholder
                no_content_text = get_text("no_content")
                wechat_final = f'<div style="padding: 40px; text-align: center; color: #999;"><p>{no_content_text}</p></div>'
            else:
                # Check if the content inside the div is actually empty (just empty div tags)
                import re as re_check
                # Remove all HTML tags and check if there's any text content
                text_content = re_check.sub(r'<[^>]+>', '', wechat_final).strip()
                has_image = "<img" in wechat_final.lower()
                if (not text_content or text_content == "") and not has_image:
                    # Empty content (and no images) - show placeholder
                    no_content_text = get_text("no_content")
                    wechat_final = f'<div style="padding: 40px; text-align: center; color: #999;"><p>{no_content_text}</p></div>'
            # Always show mobile frame when view is "Mobile"
            # Get theme values from active_theme (see themes.py for structure)
            bg_color = active_theme.get('bg', '#FAFAFA')
            card_color = active_theme.get('card', '#FFFFFF')
            text_color = active_theme.get('text', '#1D1D1F')
            muted_color = active_theme.get('muted', '#6B7280')
            primary_color = active_theme.get('primary', '#007AFF')
            accent_color = active_theme.get('accent', '#E5E5EA')
            font_family = active_theme.get('font', "-apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif")
            border_radius = active_theme.get('radius', '16px')
            shadow = active_theme.get('shadow', '0 1px 3px rgba(0,0,0,0.06)')
            
            # Mobile frame dimensions - iPhone 17 Pro style
            # Use session state which always has English values
            is_mobile = (st.session_state.get("preview_view", "Mobile") == "Mobile")
            
            # Build complete HTML with embedded styles
            preview_html = f"""
<!DOCTYPE html>
<html>
<head>
<style>
* {{ box-sizing: border-box; margin: 0; padding: 0; }}
body, html {{ 
    margin: 0; 
    padding: 0; 
    background: #f0f0f0;
    font-family: {font_family};
}}
::-webkit-scrollbar {{ display: none; }}

/* iPhone 17 Pro Frame */
.iphone-frame {{
    width: 390px;
    height: 780px;
    margin: 10px auto;
    position: relative;
    background: linear-gradient(145deg, #2a2a2e 0%, #1a1a1e 50%, #0a0a0e 100%);
    border-radius: 58px;
    padding: 12px;
    box-shadow: 
        0 50px 100px -20px rgba(0,0,0,0.5),
        0 30px 60px -30px rgba(0,0,0,0.6),
        inset 0 2px 4px rgba(255,255,255,0.1),
        inset 0 -2px 4px rgba(0,0,0,0.3),
        0 0 0 1px rgba(255,255,255,0.05);
}}
/* Dynamic Island only (no peninsula) */
.iphone-frame::after {{
    content: '';
    position: absolute;
    top: 22px;
    left: 50%;
    transform: translateX(-50%);
    width: 126px;
    height: 36px;
    background: #000;
    border-radius: 20px;
    z-index: 201;
    box-shadow: inset 0 0 3px rgba(255,255,255,0.05);
}}
/* Screen bezel */
.screen-bezel {{
    width: 100%;
    height: 100%;
    background: {bg_color};
    border-radius: 48px;
    overflow: hidden;
    position: relative;
}}
/* Side buttons - Volume */
.iphone-frame .vol-up,
.iphone-frame .vol-down {{
    position: absolute;
    left: -3px;
    width: 4px;
    background: linear-gradient(90deg, #3a3a3e, #2a2a2e);
    border-radius: 2px 0 0 2px;
}}
.iphone-frame .vol-up {{ top: 140px; height: 35px; }}
.iphone-frame .vol-down {{ top: 185px; height: 35px; }}
/* Side buttons - Power */
.iphone-frame .power {{
    position: absolute;
    right: -3px;
    top: 170px;
    width: 4px;
    height: 70px;
    background: linear-gradient(90deg, #2a2a2e, #3a3a3e);
    border-radius: 0 2px 2px 0;
}}
/* Silent switch */
.iphone-frame .silent {{
    position: absolute;
    left: -3px;
    top: 100px;
    width: 4px;
    height: 25px;
    background: linear-gradient(90deg, #3a3a3e, #2a2a2e);
    border-radius: 2px 0 0 2px;
}}

/* PC View */
.pc-frame {{
    width: 100%;
    height: 750px;
    border: 1px solid #ddd;
    border-radius: 12px;
    background: {bg_color};
    overflow: hidden;
}}

.preview-content {{
    width: 100%;
    height: 100%;
    overflow-y: auto;
    overflow-x: hidden;
    -webkit-overflow-scrolling: touch;
    padding-top: {'50px' if is_mobile else '0'};
    background: {bg_color};
    -ms-overflow-style: none;
    scrollbar-width: none;
}}

/* Canvas - theme defaults */
.mp-canvas {{ 
    padding: 16px !important;
    min-height: 100%;
    background: {bg_color};
    color: {text_color};
    font-family: {font_family};
    width: 100% !important;
    max-width: 100% !important;
    box-sizing: border-box !important;
    word-wrap: break-word;
    line-height: 1.75;
    overflow-x: hidden !important;
}}

/* Default text styles (NO !important - component inline styles override these) */
.mp-canvas p {{ 
    color: {text_color}; 
    font-family: {font_family};
    line-height: 1.75;
    margin-bottom: 16px;
}}
.mp-canvas li {{ 
    color: {text_color}; 
    font-family: {font_family};
    line-height: 1.75;
}}
.mp-canvas h1 {{ 
    color: {primary_color}; 
    font-family: {font_family};
    font-size: 24px;
    font-weight: bold;
    margin: 30px 0 20px 0;
}}
.mp-canvas h2 {{ 
    color: {primary_color}; 
    font-family: {font_family};
    font-size: 18px;
    font-weight: bold;
    margin: 30px 0 15px 0;
}}
.mp-canvas h3 {{ 
    color: {text_color}; 
    font-family: {font_family};
    font-size: 17px;
    font-weight: bold;
    margin: 20px 0 10px 0;
}}
.mp-canvas strong {{ 
    color: {primary_color}; 
    font-weight: bold;
}}
.mp-canvas blockquote {{
    border-left: 4px solid {primary_color};
    padding-left: 15px;
    color: {muted_color};
    font-style: italic;
    background: {accent_color};
    padding: 10px 10px 10px 15px;
    margin: 20px 0;
    border-radius: 0 {border_radius} {border_radius} 0;
}}
.mp-canvas a {{
    color: {primary_color};
    text-decoration: none;
}}
.mp-canvas a:hover {{
    text-decoration: underline;
}}

/* Image constraints - KEEP !important to prevent overflow */
.mp-canvas img, img {{
    max-width: 100% !important;
    height: auto !important;
    display: block;
    object-fit: contain;
    margin: 12px 0;
    border-radius: {border_radius};
}}

/* Container basics */
.mp-canvas > div {{
    max-width: 100%;
    box-sizing: border-box;
}}

/* Component wrappers - ensure isolation */
.mp-canvas .mp-steps-wrapper,
.mp-canvas .mp-timeline-wrapper {{
    display: block !important;
    clear: both !important;
    width: 100% !important;
    position: relative !important;
    overflow: visible !important;
}}

/* Steps component */
.mp-canvas .mp-steps-wrapper div[style*="display: flex"] {{
    display: flex !important;
    align-items: center !important;
    width: 100% !important;
}}
.mp-canvas span[style*="border-radius: 50%"] {{
    display: inline-flex !important;
    align-items: center !important;
    justify-content: center !important;
    flex-shrink: 0 !important;
}}

/* Timeline component */
.mp-canvas .mp-timeline-wrapper div[style*="position: relative"] {{
    position: relative !important;
    display: block !important;
    width: 100% !important;
}}
.mp-canvas .mp-timeline-wrapper span[style*="position: absolute"] {{
    position: absolute !important;
    display: block !important;
}}
.mp-canvas .mp-timeline-wrapper div[style*="border-left"] {{
    border-left-style: solid !important;
    border-left-width: 2px !important;
}}

/* Grid layout (col-2, col-3) */
.mp-canvas section[style*="display: flex"] {{
    display: flex !important;
    gap: 10px !important;
    flex-wrap: wrap !important;
}}
.mp-canvas section[style*="display: flex"] > div {{
    flex: 1 1 auto !important;
    min-width: 100px !important;
}}

/* Hero/Card - constrain width, let backgrounds show */
.mp-canvas section,
.mp-canvas > div {{
    max-width: 100% !important;
    box-sizing: border-box !important;
}}
</style>
</head>
<body>
{'<div class="iphone-frame"><div class="vol-up"></div><div class="vol-down"></div><div class="power"></div><div class="silent"></div><div class="screen-bezel">' if is_mobile else '<div class="pc-frame">'}
<div class="preview-content">
<div class="mp-canvas">{wechat_final}</div>
</div>
{'</div></div>' if is_mobile else '</div>'}
</body>
</html>
<!-- {{k}} -->"""
            st.components.v1.html(preview_html, height=850, scrolling=True)
        with t2:
            clean_code = clean_for_wechat(wechat_final)
            
            # Copy button with hidden textarea for reliable copying
            copy_button_text = get_text('copy_wechat_html')
            copy_component = f"""
            <div style="margin-bottom: 10px;">
                <textarea id="wechat-html-copy" style="position: absolute; left: -9999px; opacity: 0;">{clean_code.replace('</textarea>', '&lt;/textarea&gt;')}</textarea>
                <button onclick="copyWeChatHTML()" style="
                    background-color: #007AFF;
                    color: white;
                    border: none;
                    padding: 8px 16px;
                    border-radius: 4px;
                    cursor: pointer;
                    font-size: 14px;
                    font-weight: 500;
                    margin-bottom: 10px;
                ">{copy_button_text}</button>
                <span id="copy-feedback" style="margin-left: 10px; color: green; font-size: 14px;"></span>
            </div>
            <script>
            function copyWeChatHTML() {{
                const textarea = document.getElementById('wechat-html-copy');
                textarea.select();
                textarea.setSelectionRange(0, 99999); // For mobile devices
                
                try {{
                    const successful = document.execCommand('copy');
                    if (successful) {{
                        const feedback = document.getElementById('copy-feedback');
                        feedback.textContent = '✅ Copied!';
                        feedback.style.color = 'green';
                        setTimeout(function() {{
                            feedback.textContent = '';
                        }}, 2000);
                    }} else {{
                        // Fallback to modern API
                        navigator.clipboard.writeText(textarea.value).then(function() {{
                            const feedback = document.getElementById('copy-feedback');
                            feedback.textContent = '✅ Copied!';
                            feedback.style.color = 'green';
                            setTimeout(function() {{
                                feedback.textContent = '';
                            }}, 2000);
                        }}).catch(function(err) {{
                            alert('Failed to copy. Please select and copy manually.');
                        }});
                    }}
                }} catch (err) {{
                    // Fallback to modern API
                    navigator.clipboard.writeText(textarea.value).then(function() {{
                        const feedback = document.getElementById('copy-feedback');
                        feedback.textContent = '✅ Copied!';
                        feedback.style.color = 'green';
                        setTimeout(function() {{
                            feedback.textContent = '';
                        }}, 2000);
                    }}).catch(function(err) {{
                        alert('Failed to copy. Please select and copy manually.');
                    }});
                }}
            }}
            </script>
            """
            st.components.v1.html(copy_component, height=60)
            st.code(clean_code, language="html")
        with t3:
            col_pdf, col_word, col_html = st.columns(3)
            with col_pdf:
                # PDF Export
                if HAS_WEASYPRINT or HAS_PDFKIT or HAS_XHTML2PDF or HAS_REPORTLAB:
                    if st.button(get_text("export_pdf"), use_container_width=True, key="export_pdf_btn"):
                        with st.spinner("Generating PDF..."):
                            # Use markdown directly for cleaner PDF generation
                            pdf_html = markdown.markdown(parsed_md, extensions=['nl2br', 'extra'])
                            pdf_bytes, status = generate_pdf(pdf_html, active_theme, markdown_source=parsed_md, img_provider=img_provider)
                            
                            if pdf_bytes:
                                st.download_button(
                                    get_text("download_pdf"),
                                    pdf_bytes,
                                    "article.pdf",
                                    "application/pdf",
                                    key="download_pdf"
                                )
                                st.success("✅ PDF generated! Click download to save.")
                            else:
                                # Show user-friendly error message
                                if ErrorHandler:
                                    ErrorHandler.show_error_with_details(status)
                                else:
                                    st.error(f"❌ {status}")
                else:
                    st.info("💡 PDF: `pip install reportlab`")
            
            with col_word:
                # Word Export
                if HAS_DOCX:
                    if st.button(get_text("export_word"), use_container_width=True, key="export_word_btn"):
                        with st.spinner("Generating Word document..."):
                            word_bytes, status = generate_word(parsed_md, active_theme)
                            
                            if word_bytes:
                                st.download_button(
                                    get_text("download_word"),
                                    word_bytes,
                                    "article.docx",
                                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key="download_word"
                                )
                                st.success("✅ Word document generated!")
                            else:
                                st.error(f"❌ {status}")
                else:
                    st.info("💡 Word: `pip install python-docx`")
            
            with col_html:
                st.download_button(get_text("download_html"), standard_full, "article.html", "text/html", use_container_width=True)
            
            st.code(standard_full, language="html")

if __name__ == "__main__":
    main()
